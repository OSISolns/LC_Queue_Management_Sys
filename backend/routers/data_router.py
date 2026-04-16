"""
data_router.py
--------------
Read-heavy data endpoints:
  stats, queue aggregates, history, reports (summary + DOCX export),
  patient CRUD (read), departments/rooms/roles lists,
  settings (read), SMS history, Sukraa SOAP search pass-through.
"""
__author__ = "Valery Structure"
import logging
from datetime import datetime
from io import BytesIO
from typing import List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from jose import JWTError, jwt
from sqlalchemy import func, desc
from sqlalchemy.orm import Session, joinedload

from .. import models, schemas
from ..auth_utils import SECRET_KEY, ALGORITHM
from ..dependencies import (
    get_db, get_current_user_optional, get_current_active_user,
    get_admin_user, get_sms_officer_user, oauth2_scheme,
)
from ..services.sukraa_soap import SukraaSOAPClient

logger = logging.getLogger(__name__)
router = APIRouter(tags=["Data"])

# Sukraa SOAP Client (patient/inventory/staff autocomplete)
sukraa_soap_client = SukraaSOAPClient()


# ─── Patients ─────────────────────────────────────────────────────────────────

@router.get("/patients-all", response_model=List[schemas.PatientResponse])
def get_all_patients(skip: int = 0, limit: int = 1000000, db: Session = Depends(get_db)):
    """List all registered patients"""
    return db.query(models.Patient).offset(skip).limit(limit).all()


@router.get("/patients/search", response_model=List[schemas.PatientResponse])
def search_patients(q: str, limit: int = 5, db: Session = Depends(get_db)):
    """Search patients by name or MRN"""
    return db.query(models.Patient).filter(
        (models.Patient.first_name.ilike(f"%{q}%")) |
        (models.Patient.last_name.ilike(f"%{q}%")) |
        (models.Patient.mrn.ilike(f"%{q}%"))
    ).limit(limit).all()


@router.get("/patients/{patient_id}", response_model=schemas.PatientResponse)
def get_patient_detail(patient_id: int, db: Session = Depends(get_db)):
    """Get full patient details by ID"""
    patient = db.query(models.Patient).filter(models.Patient.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    return patient


# ─── Statistics ───────────────────────────────────────────────────────────────

@router.get("/stats")
def get_statistics(db: Session = Depends(get_db)):
    total_waiting = db.query(models.Queue).filter(models.Queue.status == "waiting").count()
    total_calling = db.query(models.Queue).filter(
        models.Queue.status.in_(["calling", "serving", "in-consultation"])
    ).count()
    total_completed = db.query(models.Queue).filter(models.Queue.status == "completed").count()
    total_no_show = db.query(models.Queue).filter(models.Queue.status == "no-show").count()

    avg_wait = 0
    completed_today = db.query(models.Queue).filter(
        models.Queue.status == "completed",
        models.Queue.completed_at >= datetime.utcnow().date()
    ).all()
    if completed_today:
        avg_wait = sum(p.wait_duration for p in completed_today) / len(completed_today)

    return {
        "total_waiting": total_waiting,
        "total_calling": total_calling,
        "total_completed": total_completed,
        "total_no_show": total_no_show,
        "avg_wait_time": round(avg_wait, 1),
        "priority_breakdown": {
            "emergency": db.query(models.Queue).filter(
                models.Queue.status == "waiting", models.Queue.priority_id == 1).count(),
            "vip": db.query(models.Queue).filter(
                models.Queue.status == "waiting", models.Queue.priority_id == 2).count(),
            "standard": db.query(models.Queue).filter(
                models.Queue.status == "waiting", models.Queue.priority_id == 3).count(),
        }
    }


@router.get("/queue/by-department")
def get_queue_by_department(db: Session = Depends(get_db)):
    results = db.query(
        models.Queue.target_dept,
        func.count(models.Queue.id).label("count")
    ).filter(models.Queue.status == "waiting").group_by(models.Queue.target_dept).all()
    return [{"department": dept, "count": count} for dept, count in results]


@router.get("/queue/by-room")
def get_queue_by_room(db: Session = Depends(get_db)):
    results = db.query(
        models.Queue.target_room,
        func.count(models.Queue.id).label("count")
    ).filter(models.Queue.status == "waiting").group_by(models.Queue.target_room).all()
    return [{"room": room, "count": count} for room, count in results]


# ─── History ──────────────────────────────────────────────────────────────────

@router.get("/history", response_model=List[schemas.QueueResponse])
def get_history(
    status: Optional[str] = None,
    department_id: Optional[int] = None,
    doctor_id: Optional[int] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    limit: int = 1000,
    db: Session = Depends(get_db),
    current_user: Optional[models.User] = Depends(get_current_user_optional)
):
    query = db.query(models.Queue)
    if status:
        query = query.filter(models.Queue.status == status)
    if department_id:
        query = query.filter(models.Queue.department_id == department_id)
    if doctor_id:
        query = query.filter(models.Queue.doctor_id == doctor_id)
    if start_date:
        query = query.filter(models.Queue.created_at >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(models.Queue.created_at <= datetime.fromisoformat(end_date))

    if current_user and current_user.role.category in ["Doctor", "Technician"]:
        query = query.filter(models.Queue.doctor_id == current_user.id)

    return query.options(
        joinedload(models.Queue.doctor),
        joinedload(models.Queue.patient),
        joinedload(models.Queue.priority),
        joinedload(models.Queue.registrar)
    ).order_by(models.Queue.created_at.desc()).limit(limit).all()


# ─── Reports ──────────────────────────────────────────────────────────────────

@router.get("/reports/summary")
def get_report_summary(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_admin_user)
):
    query = db.query(models.Queue)
    if start_date:
        query = query.filter(models.Queue.created_at >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(models.Queue.created_at <= datetime.fromisoformat(end_date))

    total = query.count()
    completed = query.filter(models.Queue.status == "completed").count()
    no_show = query.filter(models.Queue.status == "no-show").count()
    expired = query.filter(models.Queue.status == "expired").count()

    avg_wait, avg_service = 0, 0
    wait_records = query.filter(models.Queue.called_at.isnot(None)).all()
    if wait_records:
        avg_wait = sum((r.called_at - r.created_at).total_seconds() for r in wait_records) / len(wait_records)

    service_records = query.filter(
        models.Queue.completed_at.isnot(None), models.Queue.called_at.isnot(None)
    ).all()
    if service_records:
        avg_service = sum((r.completed_at - r.called_at).total_seconds() for r in service_records) / len(service_records)

    return {
        "total": total, "completed": completed, "no_show": no_show, "expired": expired,
        "avg_wait_time": round(avg_wait / 60, 2),
        "avg_service_time": round(avg_service / 60, 2)
    }


@router.get("/reports/export")
def export_history_docx(
    status: Optional[str] = None,
    department_id: Optional[int] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_admin_user)
):
    query = db.query(models.Queue).options(joinedload(models.Queue.doctor))
    if status:
        query = query.filter(models.Queue.status == status)
    if department_id:
        query = query.filter(models.Queue.department_id == department_id)
    if start_date:
        query = query.filter(models.Queue.created_at >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(models.Queue.created_at <= datetime.fromisoformat(end_date))

    records = query.all()
    doc = Document()

    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    title = doc.add_heading("Legacy Clinics - Patient History Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(9)

    if status or start_date or end_date:
        filter_p = doc.add_paragraph()
        ft = "Filters applied: "
        if status: ft += f"Status: {status} | "
        if start_date: ft += f"From: {start_date} | "
        if end_date: ft += f"To: {end_date}"
        filter_p.add_run(ft).italic = True

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=11)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Token","Patient","Dept","Reg By","Status","Created","Called","Comp","Doctor","Wait(m)","Serv(m)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(8)

    for r in records:
        row = table.add_row().cells
        for i, val in enumerate([
            str(r.token_number), str(r.patient_name or "-"), str(r.target_dept or "-"),
            str(r.registrar_name or "-"), str(r.status).capitalize(),
            r.created_at.strftime("%H:%M") if r.created_at else "-",
            r.called_at.strftime("%H:%M") if r.called_at else "-",
            r.completed_at.strftime("%H:%M") if r.completed_at else "-",
            str(r.doctor_name or "-"), str(r.wait_duration), str(r.service_duration)
        ]):
            row[i].text = val
            if row[i].paragraphs[0].runs:
                row[i].paragraphs[0].runs[0].font.size = Pt(8)

    try:
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        filename = f"Patient_History_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")


# ─── Departments / Rooms / Roles ──────────────────────────────────────────────

@router.get("/departments", response_model=List[schemas.Department])
def get_departments(db: Session = Depends(get_db)):
    return db.query(models.Department).all()


@router.get("/rooms", response_model=List[schemas.Room])
def get_rooms(db: Session = Depends(get_db)):
    return db.query(models.Room).all()


@router.get("/roles", response_model=List[schemas.Role])
def get_roles(db: Session = Depends(get_db)):
    return db.query(models.Role).all()


# ─── Settings (public read) ───────────────────────────────────────────────────

@router.get("/settings", response_model=List[schemas.SettingResponse])
def get_all_settings(db: Session = Depends(get_db)):
    """Public endpoint to get system configurations (like marquee messages)"""
    return db.query(models.Setting).all()


@router.get("/settings/{key}", response_model=schemas.SettingResponse)
def get_setting(key: str, db: Session = Depends(get_db)):
    setting = db.query(models.Setting).filter(models.Setting.key == key).first()
    if not setting:
        raise HTTPException(status_code=404, detail="Setting not found")
    return setting


# ─── SMS History (read-only) ──────────────────────────────────────────────────

@router.get("/sms/patients", response_model=List[schemas.PatientResponse])
def list_sms_patients(
    q: Optional[str] = None,
    limit: int = 20,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_sms_officer_user)
):
    query = db.query(models.Patient).filter(
        models.Patient.phone_number.isnot(None),
        models.Patient.phone_number != ""
    )
    if q:
        query = query.filter(
            (models.Patient.first_name.ilike(f"%{q}%")) |
            (models.Patient.last_name.ilike(f"%{q}%")) |
            (models.Patient.mrn.ilike(f"%{q}%")) |
            (models.Patient.phone_number.ilike(f"%{q}%"))
        )
    return query.limit(limit).all()


@router.get("/sms/patient/{patient_id}", response_model=schemas.PatientResponse)
def get_sms_patient(
    patient_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_sms_officer_user)
):
    patient = db.query(models.Patient).filter(models.Patient.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    return patient


@router.post("/sms/log")
def log_manual_sms(
    log_entry: schemas.SMSSendRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_sms_officer_user)
):
    """Log a manually sent SMS message (offline/external method e.g. MTN portal)."""
    sms_log = models.SMSHistory(
        patient_id=log_entry.patient_id,
        phone_number=log_entry.phone_number,
        message_body=log_entry.message_body,
        message_type=log_entry.message_type if isinstance(log_entry.message_type, str) else (log_entry.message_type.value if log_entry.message_type else "manual"),
        sent_by_user_id=current_user.id,
        status="manual_send",
        sent_at=datetime.utcnow()
    )
    db.add(sms_log)
    db.commit()
    db.refresh(sms_log)
    return {"status": "logged", "id": sms_log.id}


@router.get("/sms/history")
def get_sms_history(
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_sms_officer_user)
):
    history = db.query(models.SMSHistory).order_by(
        desc(models.SMSHistory.sent_at)
    ).limit(limit).all()
    return [
        {
            "id": s.id,
            "patient_name": f"{s.patient.first_name} {s.patient.last_name}" if s.patient else "Unknown",
            "phone_number": s.phone_number,
            "message_body": s.message_body,
            "message_type": s.message_type,
            "status": s.status,
            "sent_at": s.sent_at,
            "sent_by": s.sender.full_name if s.sender else "System"
        }
        for s in history
    ]


@router.get("/sms/templates")
def get_sms_templates(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_sms_officer_user)
):
    return db.query(models.SMSTemplate).all()


@router.get("/sms/export")
def export_sms_history_docx(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db),
    token: str = Depends(oauth2_scheme)
):
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username = payload.get("sub")
        if username is None:
            raise HTTPException(status_code=401, detail="Could not validate credentials")
    except JWTError:
        raise HTTPException(status_code=401, detail="Could not validate credentials")

    user = db.query(models.User).filter(models.User.username == username).first()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    query = db.query(models.SMSHistory)
    if start_date:
        try:
            query = query.filter(models.SMSHistory.sent_at >= datetime.strptime(start_date, "%Y-%m-%d"))
        except ValueError:
            pass
    if end_date:
        try:
            query = query.filter(
                models.SMSHistory.sent_at <= datetime.strptime(end_date, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            )
        except ValueError:
            pass

    history = query.order_by(desc(models.SMSHistory.sent_at)).limit(500).all()

    doc = Document()
    title_text = "SMS Communication History"
    if start_date and end_date:
        title_text += f"\n({start_date} to {end_date})"
    elif start_date:
        title_text += f"\n(From {start_date})"
    elif end_date:
        title_text += f"\n(Until {end_date})"

    t = doc.add_heading(title_text, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Generated by: {user.full_name or user.username}")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Date/Time", "Patient", "Phone", "Type", "Message"]):
        hdr[i].text = h

    for sms in history:
        row = table.add_row().cells
        row[0].text = sms.sent_at.strftime("%Y-%m-%d %H:%M")
        row[1].text = f"{sms.patient.first_name} {sms.patient.last_name}" if sms.patient else "Unknown"
        row[2].text = sms.phone_number
        row[3].text = sms.message_type.replace("_", " ").capitalize() if sms.message_type else "General"
        row[4].text = sms.message_body

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"SMS_History_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ─── Sukraa SOAP Search (pass-through) ───────────────────────────────────────

@router.get("/sukraa/search-patients")
async def search_sukraa_patients(
    q: str = "",
    limit: int = 10,
    current_user: models.User = Depends(get_current_active_user)
):
    return sukraa_soap_client.get_patients(q, count=limit)


@router.get("/sukraa/search-inventory")
async def search_sukraa_inventory(
    q: str = "",
    limit: int = 50,
    current_user: models.User = Depends(get_current_active_user)
):
    return sukraa_soap_client.get_inventory_items(q, count=limit)


@router.get("/sukraa/search-staff")
async def search_sukraa_staff(
    q: str = "",
    limit: int = 50,
    current_user: models.User = Depends(get_current_active_user)
):
    return sukraa_soap_client.get_doctors(q, count=limit)
