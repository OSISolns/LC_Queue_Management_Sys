from fastapi import FastAPI, Depends, HTTPException, status, BackgroundTasks, File, UploadFile
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func, case, or_, desc
import socketio
from typing import List, Optional
from pydantic import BaseModel
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from io import BytesIO, StringIO
from fastapi.responses import StreamingResponse
from . import models, schemas, database
from .ai_client import ai_client
from . import session_manager
from .routers import files, patient_portal
from .dependencies import (
    get_db, get_current_user, get_current_user_optional, 
    get_current_active_user, get_admin_user, get_sms_officer_user,
    oauth2_scheme, oauth2_scheme_optional, SECRET_KEY, ALGORITHM
)
import edge_tts
import uuid
import shutil
import os
import json
import csv
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from passlib.context import CryptContext
from jose import JWTError, jwt
from fastapi.middleware.cors import CORSMiddleware
import hashlib
from dotenv import load_dotenv
import logging
logger = logging.getLogger(__name__)

# SQLAdmin Imports
from sqladmin import Admin, ModelView, BaseView, expose
from sqladmin.authentication import AuthenticationBackend
from starlette.requests import Request
from starlette.responses import RedirectResponse

# Database Setup
models.Base.metadata.create_all(bind=database.engine)

# Load environment variables
load_dotenv()

# Auth Configuration
SECRET_KEY = "supersecretkey_change_me_in_production"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 600

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")
oauth2_scheme_optional = OAuth2PasswordBearer(tokenUrl="login", auto_error=False)

# SMS Configuration
# Twilio integration has been completely removed.

# Socket.IO Setup
sio = socketio.AsyncServer(async_mode='asgi', cors_allowed_origins='*')
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

socket_app = socketio.ASGIApp(sio, app)
app.include_router(files.router)
app.include_router(patient_portal.router)

# --- SMS Helper ---

def send_sms_notification(phone_number: str, token_number: str, room_number: str) -> bool:
    """
    Send SMS notification to patient with token and room number.
    Twilio integration has been removed, so this is now a stub that logs the attempt.
    """
    if not phone_number:
        logging.warning("SMS sending skipped: No phone number provided")
        return False
    
    logging.info(f"SMS notification feature removed. Would have sent token {token_number} to {phone_number}")
    return False

# --- Auth Helpers ---

def _prepare_password(password: str) -> str:
    """
    Pre-hash password with SHA-256 to avoid bcrypt's 72-byte limit.
    This allows passwords of any length to work correctly.
    """
    return hashlib.sha256(password.encode('utf-8')).hexdigest()

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

# --- Dependencies ---

async def get_sms_officer_user(current_user: models.User = Depends(get_current_active_user)):
    """Allow both Admin and SMS Officer roles to access SMS features"""
    if current_user.role.category not in ["Admin", "SMS Officer"]:
        raise HTTPException(status_code=403, detail="Not authorized for SMS operations")
    return current_user

# --- Startup ---

@app.on_event("startup")
def startup_event():
    db = database.SessionLocal()
    
    # Seed Priority Levels
    priorities = [
        {"id": 1, "name": "Emergency", "weight": 0},
        {"id": 2, "name": "VIP", "weight": 1},
        {"id": 3, "name": "Standard", "weight": 2},
    ]
    for p in priorities:
        exists = db.query(models.PriorityLevel).filter_by(id=p["id"]).first()
        if not exists:
            db.add(models.PriorityLevel(**p))
    
    # Seed Roles
    roles = [
        {"id": 1, "name": "Admin", "category": "Admin"},
        {"id": 2, "name": "Doctor", "category": "Doctor"},
        {"id": 3, "name": "Helpdesk", "category": "Helpdesk"},
        {"id": 4, "name": "Technician", "category": "Technician"},
        {"id": 5, "name": "SMS Officer", "category": "SMS Officer"},
        {"id": 6, "name": "Nurse", "category": "Nurse"},
        {"id": 100, "name": "Quality", "category": "Quality"},
    ]
    for r in roles:
        role = db.query(models.Role).filter_by(id=r["id"]).first()
        if not role:
            db.add(models.Role(**r))
        elif not role.category:
            role.category = r["category"]
            db.add(role)
            
    db.commit()

    # Seed Default Admin
    admin_role = db.query(models.Role).filter_by(name="Admin").first()
    if admin_role:
        admin_user = db.query(models.User).filter_by(username="admin").first()
        if not admin_user:
            hashed_pw = get_password_hash("admin123")
            new_admin = models.User(username="admin", hashed_password=hashed_pw, role_id=admin_role.id)
            db.add(new_admin)
            db.commit()

    # Check Queue Expiration (Auto-Clear after 10 PM)
    check_and_expire_queue(db)

    # Clean up idle/stale sessions from previous server runs
    removed = session_manager.cleanup_old_sessions(db)
    if removed:
        print(f"[SESSION] Cleaned up {removed} stale session(s) on startup.")

    db.close()

def check_and_expire_queue(db: Session):
    """
    Moves 'waiting' or 'calling' patients to 'expired' if current time > 22:00 (10 PM).
    Use this on startup or periodically.
    """
    now = datetime.now()
    # If it's past 10 PM
    # if now.hour >= 22: # Disabled for testing
    if False:
        # Find active patients
        expired_count = db.query(models.Queue).filter(
            models.Queue.status.in_(["waiting", "calling"])
        ).update(
            {models.Queue.status: "expired"}, 
            synchronize_session=False
        )
        if expired_count > 0:
            print(f"[MAINTENANCE] Expired {expired_count} patients due to time constraints (10 PM).")
            db.commit()

# --- Auth Routes ---

@app.post("/test-login")
async def test_login_endpoint(form_data: schemas.LoginRequest):
    print(f"[TEST] Received username: {form_data.username}")
    print(f"[TEST] Password length: {len(form_data.password)}")
    return {"message": "Test successful", "username": form_data.username, "password_length": len(form_data.password)}

@app.post("/login", response_model=schemas.Token)
async def login_for_access_token(request: Request, form_data: schemas.LoginRequest, db: Session = Depends(get_db)):
    try:
        origin = request.headers.get("origin")
        print(f"[LOGIN] Attempting login from {origin} for user: {form_data.username}")
        
        user = db.query(models.User).options(joinedload(models.User.role)).filter(models.User.username == form_data.username).first()
        
        if not user:
            print(f"[LOGIN] User not found: {form_data.username}")
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect username or password",
                headers={"WWW-Authenticate": "Bearer"},
            )
        
        if not verify_password(form_data.password, user.hashed_password):
            print(f"[LOGIN] Password verification failed for user: {form_data.username}")
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect username or password",
                headers={"WWW-Authenticate": "Bearer"},
            )
        
        role_title = user.role.name
        role_category = user.role.category or role_title
        
        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = create_access_token(
            data={"sub": user.username, "role": role_category}, expires_delta=access_token_expires
        )

        # Register server-side session for idle timeout tracking
        session_manager.create_session(db, access_token, user.id)
        
        return {
            "access_token": access_token, 
            "token_type": "bearer", 
            "role": role_category, 
            "role_title": role_title,
            "username": user.username, 
            "room_number": user.room_number, 
            "id": user.id,
            "full_name": user.full_name,
            "salutation": user.salutation
        }
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"[LOGIN ERROR] Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Internal server error: {str(e)}"
        )

@app.post("/logout")
async def logout(
    token: str = Depends(oauth2_scheme),
    db: Session = Depends(get_db)
):
    """Immediately invalidate the server-side session so the token can no longer be used."""
    session_manager.delete_session(db, token)
    return {"message": "Logged out successfully"}

@app.get("/users", response_model=List[schemas.User])
def get_users(role_name: Optional[str] = None, db: Session = Depends(get_db)):
    """Fetch users, optionally filtered by role name (e.g., 'Nurse', 'Doctor')"""
    query = db.query(models.User)
    if role_name:
        query = query.join(models.Role).filter(models.Role.name == role_name)
    return query.all()


# --- Queue Routes ---

@app.get("/announce")
async def announce(token: str, room: Optional[str] = None, background_tasks: BackgroundTasks = BackgroundTasks()):
    if room and room.lower() not in ["null", "undefined", "none", ""]:
        text = f"Patient {token}, please proceed to room {room}"
    else:
        text = f"Patient {token}, please proceed to your designated room"

    voice = "en-US-JennyNeural" # High quality female voice - Natural
    output_file = f"temp_{uuid.uuid4()}.mp3"
    
    communicate = edge_tts.Communicate(text, voice)
    await communicate.save(output_file)
    
    background_tasks.add_task(os.remove, output_file)
    return FileResponse(output_file, media_type="audio/mpeg")


@app.get("/sync-hims")
async def sync_hims_patients(background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    """Trigger manual sync from external Sukraa HIMS in the background"""
    from .sync_sukraa import sync_patients
    background_tasks.add_task(sync_patients)
    return {"status": "success", "message": "Synchronization started in the background"}


@app.get("/patients-all", response_model=List[schemas.PatientResponse])
def get_all_patients(skip: int = 0, limit: int = 10000, db: Session = Depends(get_db)):
    """List all registered patients"""
    return db.query(models.Patient).offset(skip).limit(limit).all()

@app.get("/patients/search", response_model=List[schemas.PatientResponse])
def search_patients(q: str, limit: int = 5, db: Session = Depends(get_db)):
    """Search patients by name or MRN"""
    query = db.query(models.Patient).filter(
        (models.Patient.first_name.ilike(f"%{q}%")) |
        (models.Patient.last_name.ilike(f"%{q}%")) |
        (models.Patient.mrn.ilike(f"%{q}%"))
    )
    return query.limit(limit).all()


@app.get("/patients/{patient_id}", response_model=schemas.PatientResponse)
def get_patient_detail(patient_id: int, db: Session = Depends(get_db)):
    """Get full patient details by ID"""
    patient = db.query(models.Patient).filter(models.Patient.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    return patient


@app.post("/register", response_model=schemas.QueueResponse)
async def register_patient(
    patient: schemas.QueueCreate, 
    db: Session = Depends(get_db),
    current_user: Optional[models.User] = Depends(get_current_user_optional)
):
    # Check if within operating hours (7 AM - 10 PM)
    now = datetime.now()
    if now.hour < 7 or now.hour >= 22:
         # Log warning but allow for now as per user instruction to "clear" not "ban"
         pass

    # Handle patient record and phone number
    patient_db_id = patient.patient_id
    
    # Validation: Ensure patient is not already active in the queue under any token
    if patient_db_id:
        active_entry = db.query(models.Queue).filter(
            models.Queue.patient_id == patient_db_id,
            models.Queue.status.in_(["waiting", "calling", "in-consultation"]),
            models.Queue.completed_at == None
        ).first()

        if active_entry:
            raise HTTPException(
                status_code=400, 
                detail=f"Patient is already active in {active_entry.target_dept or 'the queue'} with Token: {active_entry.token_number}."
            )

    priority = db.query(models.PriorityLevel).filter(models.PriorityLevel.id == patient.priority_id).first()
    if not priority:
        raise HTTPException(status_code=400, detail="Invalid Priority ID")
    
    # Validation: Ensure doctor is available today
    if patient.doctor_id:
        current_day = datetime.now().strftime("%A")
        roster_entry = db.query(models.DoctorRoster).filter(
            models.DoctorRoster.doctor_id == patient.doctor_id,
            models.DoctorRoster.day_of_week == current_day
        ).first()
        
        if roster_entry and roster_entry.status in ["not_available", "on_call"]:
            formatted_status = roster_entry.status.replace('_', ' ')
            raise HTTPException(
                status_code=400,
                detail=f"Cannot assign patient: Doctor is currently {formatted_status} today."
            )
    
    # If patient ID provided, update existing patient's phone number or gender
    if patient_db_id:
        existing_patient = db.query(models.Patient).filter(models.Patient.id == patient_db_id).first()
        if existing_patient:
            if patient.phone_number:
                existing_patient.phone_number = patient.phone_number
            if patient.gender and not existing_patient.gender:
                existing_patient.gender = patient.gender
            db.commit()
    
    # If no patient ID but phone number provided, create a new patient record
    # Patient name format: "FirstName LastName" - split it
    if not patient_db_id and patient.phone_number:
        name_parts = patient.patient_name.strip().split(maxsplit=1)
        first_name = name_parts[0] if len(name_parts) > 0 else patient.patient_name
        last_name = name_parts[1] if len(name_parts) > 1 else ""
        
        # Generate MRN
        mrn_count = db.query(models.Patient).count() + 1
        new_mrn = f"MRN{mrn_count:06d}"
        
        new_patient_record = models.Patient(
            mrn=new_mrn,
            first_name=first_name,
            last_name=last_name,
            phone_number=patient.phone_number,
            gender=patient.gender
        )
        db.add(new_patient_record)
        db.commit()
        db.refresh(new_patient_record)
        patient_db_id = new_patient_record.id
    
    prefix = priority.name[0].upper()
    count = db.query(models.Queue).filter(models.Queue.priority_id == patient.priority_id).count() + 1
    token = f"{prefix}-{count:03d}"

    new_patient = models.Queue(
        token_number=token,
        patient_name=patient.patient_name,
        patient_id=patient_db_id,
        priority_id=patient.priority_id,
        target_dept=patient.target_dept,
        target_room=patient.target_room,
        doctor_id=patient.doctor_id,
        department_id=patient.department_id,
        created_by_id=current_user.id if current_user else None,
        visit_type=patient.visit_type,
        status="waiting"
    )

    # --- AI Integration: Auto-Select Doctor/Room ---
    if not new_patient.doctor_id and (new_patient.visit_type in ["consultation", "review", "Consultation", "Review"]):
        try:
            recommendation = await ai_client.get_counter_recommendation(
                service_type=new_patient.visit_type,
                priority_id=new_patient.priority_id
            )
            if recommendation and "recommended_doctor_id" in recommendation:
                new_patient.doctor_id = recommendation["recommended_doctor_id"]
                if not new_patient.target_room:
                    new_patient.target_room = recommendation["recommended_counter_id"]
                logger.info(f"AI Auto-selected Doctor ID {new_patient.doctor_id} for Token {token}")
        except Exception as e:
            logger.error(f"AI Auto-selection failed: {e}")
    # ---------------------------------------------

    db.add(new_patient)
    db.commit()
    db.refresh(new_patient)
    
    # --- AI Integration: Predict Wait Time ---
    try:
        now_dt = datetime.utcnow()
        prediction = await ai_client.get_wait_time(
            hour=now_dt.hour,
            day_of_week=now_dt.weekday(),
            service_type=patient.visit_type or "default",
            priority=patient.priority_id
        )
        
        # Save prediction to database
        from sqlalchemy import text
        db.execute(text("""
            INSERT INTO ticket_predictions (ticket_id, model_type, model_version, predicted_value_seconds, predicted_value_str, explanation_json, created_at)
            VALUES (:tid, 'wait_time', :mver, :psec, :pstr, :expl, :cat)
        """), {
            "tid": new_patient.id,
            "mver": prediction.get('model_version', 'unknown'),
            "psec": prediction.get('predicted_wait_seconds', 0),
            "pstr": f"{prediction.get('predicted_wait_minutes', 0.0):.1f} mins",
            "expl": json.dumps(prediction),
            "cat": now_dt
        })
        db.commit()
    except Exception as e:
        logger.error(f"Failed to record AI wait time prediction: {e}")
    # -----------------------------------------
    
    # Send SMS notification if phone number is provided
    if patient.phone_number:
        room_info = patient.target_room or patient.target_dept or "TBD"
        send_sms_notification(patient.phone_number, token, room_info)
    
    await sio.emit('queue_update', {'message': 'New patient registered'})
    return new_patient

@app.post("/call-next", response_model=schemas.QueueResponse)
async def call_next_patient(request: schemas.CallNextRequest, db: Session = Depends(get_db)):
    return new_patient

@app.post("/queue/next")
async def call_next_patient(request: schemas.CallNextRequest, db: Session = Depends(get_db)):
    # Logic: Status=waiting, ORDER BY priority.weight ASC, created_at ASC
    query = db.query(models.Queue).join(models.PriorityLevel)\
        .filter(models.Queue.status == "waiting")
        
    if request.room_number:
        # Check if room is already busy with a calling patient
        active_patient = db.query(models.Queue).filter(
            models.Queue.status == "calling",
            models.Queue.room_number == request.room_number
        ).first()
        if active_patient:
            raise HTTPException(status_code=400, detail="You have an active patient. Complete them first.")

        query = query.filter(models.Queue.target_room == request.room_number)

    # Filter out patients explicitly assigned to other doctors
    if request.doctor_id:
        query = query.filter(models.Queue.doctor_id.in_([request.doctor_id, None]))

    # --- Dynamic Queue Logic (3:1 Ratio if Load >= 15) ---
    total_waiting = query.count()
    next_patient = None

    if total_waiting >= 15:
        # Fetch last 3 called/completed patients for this room today
        # Note: We use called_at as the timestamp of interaction
        history = db.query(models.Queue).filter(
            models.Queue.target_room == request.room_number,
            models.Queue.status.in_(['completed', 'calling']),
            models.Queue.called_at >= datetime.utcnow().date()
        ).order_by(models.Queue.called_at.desc()).limit(3).all()

        last_3_types = [p.visit_type for p in history]
        
        # Calculate Non-Review Streak
        non_review_streak = 0
        for vt in last_3_types:
            if vt != 'Review':
                non_review_streak += 1
            else:
                break # Streak broken by a Review

        should_call_review = (non_review_streak >= 3)
        
        if should_call_review:
            # Prioritize Review Patient
            next_patient = query.filter(models.Queue.visit_type == 'Review')\
                .order_by(models.PriorityLevel.weight.asc(), models.Queue.created_at.asc()).first()
        else:
            # Prioritize Non-Review (Consultation)
            next_patient = query.filter(models.Queue.visit_type != 'Review')\
                 .order_by(models.PriorityLevel.weight.asc(), models.Queue.created_at.asc()).first()

    # Fallback / Normal Logic (if < 15 OR preferred type not found)
    if not next_patient:
        next_patient = query.order_by(models.PriorityLevel.weight.asc(), models.Queue.created_at.asc())\
            .first()

    if not next_patient:
        raise HTTPException(status_code=404, detail="No patients in waiting queue.")

    next_patient.status = "calling"
    next_patient.doctor_id = request.doctor_id
    next_patient.room_number = request.room_number or next_patient.target_room
    next_patient.called_at = datetime.utcnow()

    db.commit()
    db.refresh(next_patient)

    is_vip = next_patient.priority.name == "VIP"

    await sio.emit('call_patient', {
        'token': next_patient.token_number,
        'room': next_patient.room_number,
        'department': next_patient.target_dept, # Added for filtering
        'name': next_patient.patient_name,
        'is_vip': is_vip
    })
    
    return next_patient

@app.post("/call-specific/{patient_id}")
async def call_specific_patient(patient_id: int, request: schemas.CallNextRequest, db: Session = Depends(get_db)):
    """Call a specific patient from the waiting list, bypassing priority order."""
    patient = db.query(models.Queue).filter(models.Queue.id == patient_id, models.Queue.status == "waiting").first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found or not waiting.")
        
    # Prevent calling another doctor's specifically assigned patient
    if patient.doctor_id is not None and request.doctor_id is not None:
        if patient.doctor_id != request.doctor_id:
             raise HTTPException(status_code=403, detail="This patient is booked for another doctor.")

    if request.room_number:
         # Check if room is already busy
        active_patient = db.query(models.Queue).filter(
            models.Queue.status == "calling",
            models.Queue.room_number == request.room_number
        ).first()
        if active_patient:
             raise HTTPException(status_code=400, detail="You have an active patient. Complete them first.")

        # Enforce Room Matching
        if patient.target_room and patient.target_room != request.room_number:
             raise HTTPException(status_code=400, detail=f"Patient is waiting for Room {patient.target_room}, you are in Room {request.room_number}")
             
        # Enforce Ordering (Next in Line OR Review)
        # Check if patient is 'Review' type
        is_review = patient.visit_type and patient.visit_type.lower() == "review"
        
        # Check if patient is effectively next in line (ignoring VIPs if standard, etc - broadly just first match)
        # Ideally we check index 0 of the waiting list for this room
        next_in_line = db.query(models.Queue).filter(
            models.Queue.status == "waiting",
            models.Queue.target_room == request.room_number
        ).order_by(models.PriorityLevel.weight.asc(), models.Queue.created_at.asc()).first()
        
        is_next = next_in_line and next_in_line.id == patient_id
        
        if not (is_review or is_next):
             raise HTTPException(status_code=400, detail="You can only call the NEXT patient or a REVIEW patient.")

    patient.status = "calling"
    patient.doctor_id = request.doctor_id
    patient.room_number = request.room_number or patient.target_room
    patient.called_at = datetime.utcnow()
    
    db.commit()
    db.refresh(patient)
    
    is_vip = patient.priority.name == "VIP"
    
    await sio.emit('call_patient', {
        'token': patient.token_number,
        'room': patient.room_number,
        'department': patient.target_dept, # Added for filtering
        'name': patient.patient_name,
        'is_vip': is_vip
    })
    
    return patient

@app.get("/queue", response_model=List[schemas.QueueResponse])
def get_queue(room: Optional[str] = None, current_user: Optional[models.User] = Depends(get_current_user_optional), db: Session = Depends(get_db)):
    query = db.query(models.Queue).filter(models.Queue.status == "waiting")
    
    # If explicitly filtered by room (e.g. from frontend param)
    if room:
        query = query.filter(models.Queue.target_room == room)
        
    # Security: Restrict by room/dept first
    if current_user and current_user.role.category in ["Doctor", "Technician", "Nurse"]:
         if current_user.room_number:
             query = query.filter(models.Queue.target_room == current_user.room_number)
         elif current_user.department_id:
             # Fallback to department if no room assigned
             dept = db.query(models.Department).filter(models.Department.id == current_user.department_id).first()
             if dept:
                 query = query.filter(models.Queue.target_dept == dept.name)
                 
         # Filter by Visit Type
         if current_user.role.category == "Doctor":
             query = query.filter(models.Queue.visit_type.in_(["consultation", "review", "Consultation", "Review", None]))
             # Hide patients specifically assigned to other doctors
             query = query.filter(models.Queue.doctor_id.in_([current_user.id, None]))
         elif current_user.role.category == "Technician":
             query = query.filter(models.Queue.visit_type.in_(["procedure", "laboratory", "pharmacy", "Procedure", "Laboratory", "Pharmacy"]))
             # Hide patients specifically assigned to other staff
             query = query.filter(models.Queue.doctor_id.in_([current_user.id, None]))

    return query.join(models.PriorityLevel).options(
        joinedload(models.Queue.doctor),
        joinedload(models.Queue.patient),
        joinedload(models.Queue.priority)
    ).order_by(models.PriorityLevel.weight.asc(), models.Queue.created_at.asc()).all()

@app.post("/complete/{patient_id}")
async def complete_patient(patient_id: int, db: Session = Depends(get_db)):
    patient = db.query(models.Queue).filter(models.Queue.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    
    patient.status = "completed"
    patient.completed_at = datetime.utcnow()
    
    # Calculate duration
    duration = 0
    if patient.called_at:
        duration = int((patient.completed_at - patient.called_at).total_seconds())

    # Create VisitHistory record
    visit = models.VisitHistory(
        patient_id=patient.patient_id,
        queue_id=patient.id,
        visit_date=patient.completed_at,
        department=patient.target_dept,
        room=patient.room_number,
        doctor_id=patient.doctor_id,
        visit_type=patient.visit_type,
        status="completed",
        duration_seconds=duration
    )
    db.add(visit)
    db.commit()
    db.refresh(visit)
    
    # Link any vitals recorded during this queue session to the permanent visit history record
    from sqlalchemy import update
    db.query(models.PatientVitals).filter(
        models.PatientVitals.queue_id == patient.id
    ).update({models.PatientVitals.visit_id: visit.id}, synchronize_session=False)
    db.commit()
    
    await sio.emit('queue_update', {'message': 'Patient completed'})
    return {"message": "Patient marked as completed"}

@app.post("/no-show/{patient_id}")
async def mark_no_show(patient_id: int, db: Session = Depends(get_db)):
    patient = db.query(models.Queue).filter(models.Queue.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    
    patient.status = "no-show"
    db.commit()
    
    await sio.emit('queue_update', {'message': 'Patient marked as no-show'})
    return {"message": "Patient marked as no-show"}

@app.post("/recall/{patient_id}")
async def recall_patient(patient_id: int, db: Session = Depends(get_db)):
    patient = db.query(models.Queue).filter(models.Queue.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    
    patient.status = "calling"
    patient.called_at = datetime.utcnow()
    db.commit()
    db.refresh(patient)
    
    await sio.emit('call_patient', {
        'token': patient.token_number,
        'room': patient.room_number,
        'department': patient.target_dept,
        'name': patient.patient_name,
        'is_vip': patient.priority.name == "VIP"
    })
    
    return {"message": "Patient recalled"}

class DoctorNotesRequest(BaseModel):
    notes: str

@app.post("/queue/{queue_id}/notes")
async def save_doctor_notes(
    queue_id: int,
    body: DoctorNotesRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_active_user)
):
    """Save or update doctor's consultation notes. Only allowed while serving or after completion."""
    entry = db.query(models.Queue).filter(models.Queue.id == queue_id).first()
    if not entry:
        raise HTTPException(status_code=404, detail="Queue entry not found")
    if entry.status == "waiting":
        raise HTTPException(status_code=403, detail="Cannot add notes before the patient is being served")
    entry.doctor_notes = body.notes.strip()
    db.commit()
    return {"message": "Notes saved", "doctor_notes": entry.doctor_notes}

@app.get("/stats")
def get_statistics(db: Session = Depends(get_db)):
    total_waiting = db.query(models.Queue).filter(models.Queue.status == "waiting").count()
    total_calling = db.query(models.Queue).filter(models.Queue.status == "calling").count()
    total_completed = db.query(models.Queue).filter(models.Queue.status == "completed").count()
    total_no_show = db.query(models.Queue).filter(models.Queue.status == "no-show").count()
    
    emergency_waiting = db.query(models.Queue).filter(
        models.Queue.status == "waiting",
        models.Queue.priority_id == 1
    ).count()
    vip_waiting = db.query(models.Queue).filter(
        models.Queue.status == "waiting",
        models.Queue.priority_id == 2
    ).count()
    standard_waiting = db.query(models.Queue).filter(
        models.Queue.status == "waiting",
        models.Queue.priority_id == 3
    ).count()
    
    return {
        "total_waiting": total_waiting,
        "total_calling": total_calling,
        "total_completed": total_completed,
        "total_no_show": total_no_show,
        "priority_breakdown": {
            "emergency": emergency_waiting,
            "vip": vip_waiting,
            "standard": standard_waiting
        }
    }

@app.get("/queue/by-department")
def get_queue_by_department(db: Session = Depends(get_db)):
    from sqlalchemy import func
    results = db.query(
        models.Queue.target_dept,
        func.count(models.Queue.id).label('count')
    ).filter(
        models.Queue.status == "waiting"
    ).group_by(models.Queue.target_dept).all()
    
    return [{"department": dept, "count": count} for dept, count in results]

@app.get("/queue/by-room")
def get_queue_by_room(db: Session = Depends(get_db)):
    from sqlalchemy import func
    results = db.query(
        models.Queue.target_room,
        func.count(models.Queue.id).label('count')
    ).filter(
        models.Queue.status == "waiting"
    ).group_by(models.Queue.target_room).all()
    
    return [{"room": room, "count": count} for room, count in results]

@app.get("/history", response_model=List[schemas.QueueResponse])
def get_history(
    status: Optional[str] = None, 
    department_id: Optional[int] = None,
    doctor_id: Optional[int] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    limit: int = 1000, 
    db: Session = Depends(get_db), 
    current_user: Optional[models.User] = Depends(get_current_user_optional)
):
    query = db.query(models.Queue)
    if status:
        query = query.filter(models.Queue.status == status)
    if department_id:
        query = query.filter(models.Queue.department_id == department_id)
    if doctor_id:
        query = query.filter(models.Queue.doctor_id == doctor_id)
    if start_date:
        query = query.filter(models.Queue.created_at >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(models.Queue.created_at <= datetime.fromisoformat(end_date))
        
    # Ensure doctors/technicians/nurses only see relevant history
    if current_user and current_user.role.category in ["Doctor", "Technician"]:
        query = query.filter(models.Queue.doctor_id == current_user.id)
        
    return query.options(
        joinedload(models.Queue.doctor),
        joinedload(models.Queue.patient),
        joinedload(models.Queue.priority),
        joinedload(models.Queue.registrar)
    ).order_by(models.Queue.created_at.desc()).limit(limit).all()

@app.get("/reports/summary")
def get_report_summary(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_admin_user)
):
    query = db.query(models.Queue)
    if start_date:
        query = query.filter(models.Queue.created_at >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(models.Queue.created_at <= datetime.fromisoformat(end_date))

    total = query.count()
    completed = query.filter(models.Queue.status == "completed").count()
    no_show = query.filter(models.Queue.status == "no-show").count()
    expired = query.filter(models.Queue.status == "expired").count()

    # Wait times (created_at to called_at)
    avg_wait = 0
    wait_records = query.filter(models.Queue.called_at.isnot(None)).all()
    if wait_records:
        total_wait = sum((r.called_at - r.created_at).total_seconds() for r in wait_records)
        avg_wait = total_wait / len(wait_records)

    # Service times (called_at to completed_at)
    avg_service = 0
    service_records = query.filter(models.Queue.completed_at.isnot(None), models.Queue.called_at.isnot(None)).all()
    if service_records:
        total_service = sum((r.completed_at - r.called_at).total_seconds() for r in service_records)
        avg_service = total_service / len(service_records)

    return {
        "total": total,
        "completed": completed,
        "no_show": no_show,
        "expired": expired,
        "avg_wait_time": round(avg_wait / 60, 2), # in minutes
        "avg_service_time": round(avg_service / 60, 2) # in minutes
    }

@app.get("/reports/export")
def export_history_docx(
    status: Optional[str] = None,
    department_id: Optional[int] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_admin_user)
):
    query = db.query(models.Queue).options(joinedload(models.Queue.doctor))
    if status: query = query.filter(models.Queue.status == status)
    if department_id: query = query.filter(models.Queue.department_id == department_id)
    if start_date: query = query.filter(models.Queue.created_at >= datetime.fromisoformat(start_date))
    if end_date: query = query.filter(models.Queue.created_at <= datetime.fromisoformat(end_date))

    records = query.all()
    
    doc = Document()
    
    # Set orientation to Landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    # Title
    title = doc.add_heading('Legacy Clinics - Patient History Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    
    # Filter hint
    if status or start_date or end_date:
        filter_p = doc.add_paragraph()
        filter_text = "Filters applied: "
        if status: filter_text += f"Status: {status} | "
        if start_date: filter_text += f"From: {start_date} | "
        if end_date: filter_text += f"To: {end_date}"
        filter_p.add_run(filter_text).italic = True
    
    doc.add_paragraph() # Spacer
    
    # Table creation - 11 columns
    table = doc.add_table(rows=1, cols=11)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    headers = [
        "Token", "Patient", "Dept", "Reg By", "Status", 
        "Created", "Called", "Comp", "Doctor", "Wait(m)", "Serv(m)"
    ]
    for i, head in enumerate(headers):
        hdr_cells[i].text = head
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(8)
    
    # Data rows
    for r in records:
        row_cells = table.add_row().cells
        data = [
            str(r.token_number),
            str(r.patient_name or "-"),
            str(r.target_dept or "-"),
            str(r.registrar_name or "-"),
            str(r.status).capitalize(),
            r.created_at.strftime("%H:%M") if r.created_at else "-",
            r.called_at.strftime("%H:%M") if r.called_at else "-",
            r.completed_at.strftime("%H:%M") if r.completed_at else "-",
            str(r.doctor_name or "-"),
            str(r.wait_duration),
            str(r.service_duration)
        ]
        for i, val in enumerate(data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)

    # Save to stream
    try:
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        filename = f"Patient_History_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        print(f"[REPORTS] DOCX Export Error: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")

# --- Admin Management Routes ---

@app.post("/users", response_model=schemas.User)
def create_user(user: schemas.UserCreate, db: Session = Depends(get_db), current_user: models.User = Depends(get_admin_user)):
    db_user = db.query(models.User).filter(models.User.username == user.username).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    hashed_password = get_password_hash(user.password)
    db_user = models.User(
        username=user.username, 
        hashed_password=hashed_password, 
        role_id=user.role_id, 
        is_active=user.is_active,
        department_id=user.department_id,
        room_number=user.room_number,
        full_name=user.full_name,
        email=user.email,
        phone_number=user.phone_number,
        salutation=user.salutation
    )
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    return db_user

@app.get("/users", response_model=List[schemas.User])
def read_users(skip: int = 0, limit: int = 100, db: Session = Depends(get_db), current_user: models.User = Depends(get_current_active_user)):
    users = db.query(models.User).offset(skip).limit(limit).all()
    return users

@app.get("/public/doctors", response_model=List[schemas.User])
def get_public_doctors(db: Session = Depends(get_db)):
    """A public endpoint to fetch active doctors without authentication, primarily used by the kiosk."""
    doctor_role = db.query(models.Role).filter(models.Role.name == "Doctor").first()
    if not doctor_role:
        return []
    
    doctors = db.query(models.User).filter(
        models.User.role_id == doctor_role.id,
        models.User.is_active == True
    ).all()
    return doctors

@app.put("/users/{user_id}", response_model=schemas.User)
def update_user(user_id: int, user_update: schemas.UserUpdate, db: Session = Depends(get_db), current_user: models.User = Depends(get_admin_user)):
    db_user = db.query(models.User).filter(models.User.id == user_id).first()
    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if user_update.role_id is not None:
        db_user.role_id = user_update.role_id
    if user_update.department_id is not None:
        db_user.department_id = user_update.department_id
    if user_update.room_number is not None:
        db_user.room_number = user_update.room_number
    if user_update.is_active is not None:
        db_user.is_active = user_update.is_active
    if user_update.full_name is not None:
        db_user.full_name = user_update.full_name
    if user_update.email is not None:
        db_user.email = user_update.email
    if user_update.phone_number is not None:
        db_user.phone_number = user_update.phone_number
    if user_update.salutation is not None:
        db_user.salutation = user_update.salutation
    if user_update.password:
        db_user.hashed_password = get_password_hash(user_update.password)
        
    db.commit()
    db.refresh(db_user)
    return db_user

@app.delete("/users/{user_id}")
def delete_user(user_id: int, request: schemas.UserDeleteRequest, db: Session = Depends(get_db), current_user: models.User = Depends(get_admin_user)):
    # Prevent self-deletion
    if user_id == current_user.id:
        raise HTTPException(status_code=400, detail="Cannot delete your own account")
    
    # Verify Admin Password
    if not verify_password(request.admin_password, current_user.hashed_password):
        raise HTTPException(status_code=401, detail="Incorrect admin password")
    
    db_user = db.query(models.User).filter(models.User.id == user_id).first()
    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")
    
    db.delete(db_user)
    db.commit()
    print(f"[AUDIT] User {current_user.username} deleted user {db_user.username}. Reason: {request.reason}")
    return {"message": "User deleted successfully"}

@app.post("/departments", response_model=schemas.Department)
def create_department(dept: schemas.DepartmentCreate, db: Session = Depends(get_db), current_user: models.User = Depends(get_admin_user)):
    db_dept = models.Department(name=dept.name)
    db.add(db_dept)
    db.commit()
    db.refresh(db_dept)
    return db_dept

@app.get("/departments", response_model=List[schemas.Department])
def read_departments(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    return db.query(models.Department).offset(skip).limit(limit).all()

@app.post("/rooms", response_model=schemas.Room)
def create_room(room: schemas.RoomCreate, db: Session = Depends(get_db), current_user: models.User = Depends(get_admin_user)):
    db_room = models.Room(name=room.name, department_id=room.department_id, floor=room.floor)
    db.add(db_room)
    db.commit()
    db.refresh(db_room)
    return db_room

@app.put("/rooms/{room_id}", response_model=schemas.Room)
def update_room(room_id: int, room: schemas.RoomUpdate, db: Session = Depends(get_db), current_user: models.User = Depends(get_admin_user)):
    db_room = db.query(models.Room).filter(models.Room.id == room_id).first()
    if not db_room:
        raise HTTPException(status_code=404, detail="Room not found")
        
    update_data = room.dict(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_room, key, value)
        
    db.commit()
    db.refresh(db_room)
    return db_room

@app.get("/rooms", response_model=List[schemas.Room])
def read_rooms(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    return db.query(models.Room).offset(skip).limit(limit).all()

@app.get("/roles", response_model=List[schemas.Role])
def read_roles(db: Session = Depends(get_db)):
    return db.query(models.Role).all()

# --- SMS Routes ---

@app.get("/sms/patients", response_model=List[schemas.PatientForSMS])
def get_patients_for_sms(
    q: Optional[str] = None, 
    skip: int = 0, 
    limit: int = 50, 
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_sms_officer_user)
):
    """Get patients with phone numbers for SMS communications"""
    query = db.query(
        models.Patient.id,
        models.Patient.mrn,
        models.Patient.first_name,
        models.Patient.last_name,
        models.Patient.phone_number,
        func.max(models.VisitHistory.visit_date).label('last_visit_date')
    ).outerjoin(
        models.VisitHistory, models.Patient.id == models.VisitHistory.patient_id
    ).filter(
        models.Patient.phone_number.isnot(None),
        models.Patient.phone_number != ""
    ).group_by(
        models.Patient.id,
        models.Patient.mrn,
        models.Patient.first_name,
        models.Patient.last_name,
        models.Patient.phone_number
    )
    
    # Search filter
    if q:
        search_pattern = f"%{q}%"
        query = query.filter(
            or_(
                models.Patient.first_name.ilike(search_pattern),
                models.Patient.last_name.ilike(search_pattern),
                models.Patient.mrn.ilike(search_pattern),
                models.Patient.phone_number.ilike(search_pattern)
            )
        )
    
    results = query.order_by(models.Patient.created_at.desc()).offset(skip).limit(limit).all()
    
    # Convert to schema format
    return [
        schemas.PatientForSMS(
            id=r.id,
            mrn=r.mrn,
            first_name=r.first_name,
            last_name=r.last_name,
            phone_number=r.phone_number,
            last_visit_date=r.last_visit_date
        ) for r in results
    ]

@app.get("/sms/patient/{patient_id}")
def get_patient_for_sms(
    patient_id: int,
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_sms_officer_user)
):
    """Get limited patient details for SMS (no medical information)"""
    patient = db.query(models.Patient).filter(models.Patient.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
    
    # Get visit count
    visit_count = db.query(models.VisitHistory).filter(
        models.VisitHistory.patient_id == patient_id
    ).count()
    
    # Get last visit date
    last_visit = db.query(models.VisitHistory).filter(
        models.VisitHistory.patient_id == patient_id
    ).order_by(models.VisitHistory.visit_date.desc()).first()
    
    # Get upcoming queue entries
    upcoming = db.query(models.Queue).filter(
        models.Queue.patient_id == patient_id,
        models.Queue.status == "waiting"
    ).all()
    
    return {
        "id": patient.id,
        "mrn": patient.mrn,
        "first_name": patient.first_name,
        "last_name": patient.last_name,
        "phone_number": patient.phone_number,
        "email": patient.email,
        "visit_count": visit_count,
        "last_visit_date": last_visit.visit_date if last_visit else None,
        "has_upcoming_appointment": len(upcoming) > 0
    }

@app.post("/sms/send", response_model=schemas.SMSHistoryResponse)
def send_sms_message(
    request: schemas.SMSSendRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_sms_officer_user)
):
    """Send SMS notification to a patient (asynchronous)"""
    # Validate phone number
    if not request.phone_number:
        raise HTTPException(status_code=400, detail="Phone number is required")
    
    # Create SMS record immediately with 'pending' status
    sms_record = models.SMSHistory(
        patient_id=request.patient_id,
        phone_number=request.phone_number,
        message_body=request.message_body,
        message_type=request.message_type.value if request.message_type else "general",
        sent_by_user_id=current_user.id,
        status="pending"
    )
    db.add(sms_record)
    db.commit()
    db.refresh(sms_record)
    
    # Send SMS in background
    background_tasks.add_task(
        send_sms_background, 
        sms_record.id, 
        request.phone_number, 
        request.message_body
    )
    
    return sms_record

def send_sms_background(sms_id: int, phone_number: str, message_body: str):
    """Background task to send SMS (Twilio removed)"""
    db = database.SessionLocal()
    try:
        sms_record = db.query(models.SMSHistory).filter(models.SMSHistory.id == sms_id).first()
        if not sms_record:
            logging.error(f"SMS record {sms_id} not found")
            return
        
        logging.info(f"SMS sending disabled (Twilio removed). Would have sent to {phone_number}: {message_body}")
        sms_record.status = "not_configured"
        
        db.commit()
    except Exception as e:
        logging.error(f"Error in background SMS task: {e}")
    finally:
        db.close()


@app.get("/sms/history", response_model=List[schemas.SMSHistoryResponse])
def get_sms_history(
    patient_id: Optional[int] = None,
    message_type: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_sms_officer_user)
):
    """Get SMS history with optional filters"""
    query = db.query(models.SMSHistory)
    
    if patient_id:
        query = query.filter(models.SMSHistory.patient_id == patient_id)
    
    if message_type:
        query = query.filter(models.SMSHistory.message_type == message_type)
        
    if start_date:
        try:
            start = datetime.strptime(start_date, "%Y-%m-%d")
            query = query.filter(models.SMSHistory.sent_at >= start)
        except ValueError:
            pass
            
    if end_date:
        try:
            end = datetime.strptime(end_date, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            query = query.filter(models.SMSHistory.sent_at <= end)
        except ValueError:
            pass
    
    return query.order_by(models.SMSHistory.sent_at.desc()).offset(skip).limit(limit).all()

@app.get("/sms/templates", response_model=List[schemas.MessageTemplate])
def get_message_templates(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_sms_officer_user)
):
    """Get predefined SMS message templates from database"""
    templates = db.query(models.SMSTemplate).all()
    # Convert to schema format
    return [
        schemas.MessageTemplate(
            type=t.type,
            template=t.template,
            description=t.description
        ) for t in templates
    ]

@app.post("/sms/log", response_model=schemas.SMSHistoryResponse)
def log_manual_sms(
    request: schemas.SMSSendRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_sms_officer_user)
):
    """Log a manually sent SMS (via MTN Portal)"""
    new_sms = models.SMSHistory(
        patient_id=request.patient_id,
        phone_number=request.phone_number,
        message_body=request.message_body,
        message_type=request.message_type,
        sent_by_user_id=current_user.id,
        status="manual_send",
        sent_at=datetime.utcnow()
    )
    db.add(new_sms)
    db.commit()
    db.refresh(new_sms)
    db.commit()
    db.refresh(new_sms)
    return new_sms

@app.get("/sms/export")
def export_sms_history(
    token: str,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """Export SMS history to Word document"""
    try:
        # Manual Authentication via Query Param
        if not token:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        try:
            # Check for Bearer prefix and strip it if present
            if token.startswith("Bearer "):
                token = token.split(" ")[1]

            payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
            username: str = payload.get("sub")
            if username is None:
                 raise HTTPException(status_code=401, detail="Could not validate credentials")
        except JWTError:
            raise HTTPException(status_code=401, detail="Could not validate credentials")
            
        user = db.query(models.User).filter(models.User.username == username).first()
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
            
        current_user = user

        # Fetch all history, ordered by newest first
        query = db.query(models.SMSHistory)
        
        if start_date:
            try:
                start = datetime.strptime(start_date, "%Y-%m-%d")
                query = query.filter(models.SMSHistory.sent_at >= start)
            except ValueError:
                pass
                
        if end_date:
            try:
                end = datetime.strptime(end_date, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
                query = query.filter(models.SMSHistory.sent_at <= end)
            except ValueError:
                pass
        
        history = query.order_by(desc(models.SMSHistory.sent_at)).limit(500).all()
        
        # Create Word Document
        doc = Document()
        
        # Title
        title_text = 'SMS Communication History'
        if start_date and end_date:
            title_text += f'\n({start_date} to {end_date})'
        elif start_date:
            title_text += f'\n(From {start_date})'
        elif end_date:
            title_text += f'\n(Until {end_date})'
            
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Generated by: {current_user.full_name or current_user.username}")
        
        # Table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date/Time'
        hdr_cells[1].text = 'Patient'
        hdr_cells[2].text = 'Phone'
        hdr_cells[3].text = 'Type'
        hdr_cells[4].text = 'Message'
        
        # Data Rows
        for sms in history:
            row_cells = table.add_row().cells
            row_cells[0].text = sms.sent_at.strftime('%Y-%m-%d %H:%M')
            row_cells[1].text = f"{sms.patient.first_name} {sms.patient.last_name}" if sms.patient else "Unknown"
            row_cells[2].text = sms.phone_number
            row_cells[3].text = sms.message_type.replace('_', ' ').capitalize() if sms.message_type else "General"
            row_cells[4].text = sms.message_body
            
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"SMS_History_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        
        return StreamingResponse(
            buffer, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/panel-links")
def get_panel_links(request: Request, current_user: models.User = Depends(get_current_active_user)):
    # Dynamically determine base_url from request (assuming frontend is on same host, port 5173)
    host = request.url.hostname
    base_url = f"http://{host}:5173"
    
    return [
        {"name": "Admin Dashboard", "url": f"{base_url}/admin", "role_required": "Admin"},
        {"name": "Doctor Dashboard", "url": f"{base_url}/doctor", "role_required": "Doctor"},
        {"name": "Kiosk Registration", "url": f"{base_url}/kiosk", "role_required": "Helpdesk"},
        {"name": "Public Display", "url": f"{base_url}/display", "role_required": "None"},
    ]

@app.post("/reset-queue")
def reset_queue(db: Session = Depends(get_db), current_user: models.User = Depends(get_admin_user)):
    """Clear all waiting patients"""
    db.query(models.Queue).filter(models.Queue.status == "waiting").delete()
    db.commit()
    return {"message": "Queue cleared"}

@app.post("/reset-calling")
async def reset_calling(db: Session = Depends(get_db), current_user: models.User = Depends(get_admin_user)):
    """Clear all calling patients (mark as expired to remove from display)"""
    num_updated = db.query(models.Queue).filter(models.Queue.status == "calling").update(
        {models.Queue.status: "expired"}, 
        synchronize_session=False
    )
    db.commit()
    print(f"[ADMIN] Cleared {num_updated} active calls via reset-calling")
    await sio.emit('queue_update', {'message': 'Active calls cleared'})
    return {"message": f"Active calls cleared: {num_updated}"}

@app.delete("/history")
def delete_history(
    request: schemas.HistoryDeleteRequest,
    db: Session = Depends(get_db), 
    current_user: models.User = Depends(get_admin_user)
):
    """
    Delete history records (completed, no-show, expired) within a date range.
    Requires admin password and reason.
    """
    # Verify Admin Password
    if not verify_password(request.admin_password, current_user.hashed_password):
        raise HTTPException(status_code=401, detail="Incorrect admin password")
    
    if not request.reason or len(request.reason) < 5:
        raise HTTPException(status_code=400, detail="A valid reason (min 5 chars) is required for deletion")

    query = db.query(models.Queue).filter(models.Queue.status.in_(["completed", "no-show", "expired"]))
    
    if request.start_date:
        query = query.filter(models.Queue.created_at >= request.start_date)
    if request.end_date:
        query = query.filter(models.Queue.created_at <= request.end_date)
        
    deleted_count = query.delete(synchronize_session=False)
    db.commit()
    
    print(f"[AUDIT] User {current_user.username} purged {deleted_count} history records. Reason: {request.reason}")
    
    return {"message": f"Successfully purged {deleted_count} records"}

# --- Settings API ---
@app.get("/settings", response_model=List[schemas.SettingResponse])
def get_all_settings(db: Session = Depends(get_db)):
    """Public endpoint to get system configurations (like marquee messages)"""
    return db.query(models.Setting).all()

@app.get("/settings/{key}", response_model=schemas.SettingResponse)
def get_setting(key: str, db: Session = Depends(get_db)):
    setting = db.query(models.Setting).filter(models.Setting.key == key).first()
    if not setting:
        raise HTTPException(status_code=404, detail="Setting not found")
    return setting

@app.post("/settings", response_model=schemas.SettingResponse)
def update_setting(
    setting_in: schemas.SettingCreate, 
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_admin_user)
):
    """Admin endpoint to create or update a setting"""
    setting = db.query(models.Setting).filter(models.Setting.key == setting_in.key).first()
    if setting:
        setting.value = setting_in.value
        setting.description = setting_in.description or setting.description
    else:
        setting = models.Setting(**setting_in.dict())
        db.add(setting)
    
    db.commit()
    db.refresh(setting)
    return setting

@sio.event
async def connect(sid, environ):
    print("connect ", sid)


# --- Admin Panel Configuration (GUI) ---

class AdminAuth(AuthenticationBackend):
    async def login(self, request: Request) -> bool:
        form = await request.form()
        username, password = form.get("username"), form.get("password")
        
        # Use a fresh session for auth check
        with database.SessionLocal() as db:
            user = db.query(models.User).filter(models.User.username == username).first()
            if not user:
                return False
            # Verify password using the helper from main.py
            if not verify_password(password, user.hashed_password):
                return False
            # Ensure only Admin can access GUI
            if user.role.name != "Admin":
                return False
        
        # Determine strict session token
        request.session.update({"token": username})
        return True

    async def logout(self, request: Request) -> bool:
        request.session.clear()
        return True

    async def authenticate(self, request: Request) -> bool:
        token = request.session.get("token")
        if not token:
            return False
        return True

# Admin Views

from wtforms import PasswordField, SelectField

class UserAdmin(ModelView, model=models.User):
    column_list = [models.User.id, models.User.username, models.User.salutation, models.User.full_name, "is_available", models.User.role, models.User.last_login]
    column_sortable_list = [models.User.id, models.User.username, models.User.full_name, models.User.email, models.User.last_login]
    column_searchable_list = [models.User.username, models.User.full_name, models.User.email]
    icon = "fa-solid fa-user"
    can_create = True
    can_edit = True
    can_delete = True
    can_view_details = True
    can_export = True
    category = "User Management"

    form_columns = [
        models.User.username,
        models.User.salutation,
        models.User.full_name,
        models.User.email,
        models.User.phone_number,
        models.User.role,
        models.User.department,
        models.User.room_number,
        models.User.is_active,
        models.User.is_available
    ]
    
    form_extra_fields = {
        "password": PasswordField("Password"),
        "salutation": SelectField("Salutation", choices=[
            ("", "Select..."),
            ("Dr.", "Dr."),
            ("Mr.", "Mr."),
            ("Mrs.", "Mrs."),
            ("Ms.", "Ms.")
        ])
    }

    async def on_model_change(self, data, model, is_created, request):
        password = data.get("password")
        if password:
            model.hashed_password = get_password_hash(password)
        elif is_created:
            # Require password for new users if not handled by form validators (omitted for flexibility here)
            pass 
            
        # Clean up transient field
        if "password" in data:
            del data["password"]

class RoleAdmin(ModelView, model=models.Role):
    column_list = [models.Role.id, models.Role.name]
    can_delete = False # Protect roles
    icon = "fa-solid fa-user-shield"
    can_view_details = True
    category = "User Management"

class DepartmentAdmin(ModelView, model=models.Department):
    column_list = [models.Department.id, models.Department.name]
    icon = "fa-solid fa-building"
    can_export = True
    category = "Hospital Resources"

class RoomAdmin(ModelView, model=models.Room):
    column_list = [models.Room.id, models.Room.name, models.Room.department, models.Room.floor, models.Room.extension]
    column_sortable_list = [models.Room.id, models.Room.name, models.Room.floor]
    column_searchable_list = [models.Room.name, models.Room.extension]
    icon = "fa-solid fa-door-open"
    can_export = True
    category = "Hospital Resources"

    form_overrides = {"floor": SelectField}
    form_args = {"floor": {"choices": [("", "None"), ("ground", "Ground Floor"), ("first", "First Floor"), ("pediatrics", "Pediatrics")]}}

class PatientAdmin(ModelView, model=models.Patient):
    column_list = [models.Patient.mrn, models.Patient.first_name, models.Patient.last_name, models.Patient.date_of_birth, models.Patient.is_active]
    column_sortable_list = [models.Patient.mrn, models.Patient.last_name, models.Patient.date_of_birth]
    column_searchable_list = [models.Patient.mrn, models.Patient.last_name, models.Patient.first_name]
    icon = "fa-solid fa-hospital-user"
    can_export = True
    can_view_details = True
    category = "Medical Records"

class QueueAdmin(ModelView, model=models.Queue):
    column_list = [models.Queue.token_number, models.Queue.patient_name, models.Queue.priority, models.Queue.status, models.Queue.target_dept, models.Queue.doctor, models.Queue.room_number, models.Queue.doctor_notes, "duration"]
    column_sortable_list = [models.Queue.created_at]
    icon = "fa-solid fa-list-ol"
    can_view_details = True
    category = "Queue Management"

    def duration(self, model):
        if model.called_at and model.completed_at:
            diff = model.completed_at - model.called_at
            # Format as MM:SS or HH:MM:SS
            total_seconds = int(diff.total_seconds())
            minutes = total_seconds // 60
            seconds = total_seconds % 60
            return f"{minutes}m {seconds}s"
        return "-"

class VisitHistoryAdmin(ModelView, model=models.VisitHistory):
    column_list = [models.VisitHistory.visit_date, models.VisitHistory.patient, models.VisitHistory.doctor, models.VisitHistory.status, models.VisitHistory.doctor_notes]
    column_sortable_list = [models.VisitHistory.visit_date, models.VisitHistory.status]
    icon = "fa-solid fa-file-medical"
    can_view_details = True
    category = "Medical Records"

class PatientVitalsAdmin(ModelView, model=models.PatientVitals):
    name = "Triage / Vitals"
    icon = "fa-solid fa-heart-pulse"
    category = "Medical Records"
    column_list = [models.PatientVitals.id, models.PatientVitals.patient_id, models.PatientVitals.temperature, models.PatientVitals.blood_pressure, models.PatientVitals.recorded_at]

class PriorityLevelAdmin(ModelView, model=models.PriorityLevel):
    column_list = [models.PriorityLevel.name, models.PriorityLevel.weight]
    can_delete = False
    icon = "fa-solid fa-layer-group"

from fastapi.staticfiles import StaticFiles

# Mount Static Files
app.mount("/static", StaticFiles(directory="backend/static"), name="static")

# Custom Admin Class for Dashboard
class LCAdmin(Admin):
    async def index(self, request: Request):
        # Fetch stats for dashboard
        with database.SessionLocal() as db:
            waiting_count = db.query(models.Queue).filter(models.Queue.status == 'waiting').count()
            active_visits = db.query(models.Queue).filter(models.Queue.status == 'serving').count() # Using Queue Serving status for active visits
            
            today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
            total_today = db.query(models.VisitHistory).filter(models.VisitHistory.visit_date >= today_start).count()
            
            # Active staff query
            active_staff = db.query(models.User).filter(models.User.is_active == True).count()

            # Extra stats — patients, SMS, files, rosters
            total_patients = db.query(models.Patient).filter(models.Patient.is_active == True).count()
            sms_today = db.query(models.SMSHistory).filter(
                models.SMSHistory.sent_at >= today_start
            ).count()
            files_count = db.query(models.Document).filter(models.Document.is_active == True).count()
            from backend.roster.models import RosterDay
            active_rosters = db.query(RosterDay).filter(RosterDay.status == "published").count()

            
            # Live Queue (Top 5: Serving first, then Priority)
            live_queue = db.query(models.Queue).filter(
                models.Queue.status.in_(['waiting', 'serving'])
            ).order_by(
                case((models.Queue.status == 'serving', 0), else_=1),
                models.Queue.priority_id.asc(),
                models.Queue.created_at.asc()
            ).limit(5).all()
            
            # Doctor Status (Real-time Availability)
            # Fetch users with medical roles
            # User.role is a relationship, so we must join Role table
            medical_staff = db.query(models.User).join(models.Role).filter(
                models.User.is_active == True,
                or_(
                    models.Role.name == 'Doctor', 
                    models.Role.name == 'Nurse',
                    models.Role.name == 'Technician'
                )
            ).all()
            
            doctor_status = []
            for staff in medical_staff:
                # Check if they have an active patient
                # We check IF they are assigned to a queue item that is currently 'serving'
                active_pt = db.query(models.Queue).filter(
                    models.Queue.doctor_id == staff.id,
                    models.Queue.status == 'serving'
                ).first()
                
                status_label = "Busy" if active_pt else "Available"
                # staff.role is an object, so we access .name
                role_name = staff.role.name.capitalize() if staff.role and staff.role.name else "Staff"
                staff_name = staff.full_name if staff.full_name else staff.username
                
                doctor_status.append({
                    "name": staff_name,
                    "role": role_name,
                    "status": status_label,
                    "current_patient": active_pt.patient_name if active_pt else None
                })
            
        return await self.templates.TemplateResponse(request, "admin/index.html", context={
            "request": request,
            "waiting_count": waiting_count,
            "active_visits": active_visits,
            "total_today": total_today,
            "active_staff": active_staff,
            "total_patients": total_patients,
            "sms_today": sms_today,
            "files_count": files_count,
            "active_rosters": active_rosters,
            "live_queue": live_queue,
            "doctor_status": doctor_status
        })




# Initialize Admin
authentication_backend = AdminAuth(secret_key=SECRET_KEY)
admin = LCAdmin(app, database.engine, authentication_backend=authentication_backend, title="Legacy Clinics Admin", logo_url="/static/logo.png", templates_dir="backend/templates")

class AnalyticsView(BaseView):
    name = "Analytics"
    icon = "fa-solid fa-chart-line"

    @expose("/analytics", methods=["GET"])
    async def analytics(self, request: Request):
        with database.SessionLocal() as db:
             # Date Filter Logic
             start_str = request.query_params.get("start_date")
             end_str = request.query_params.get("end_date")
             
             now = datetime.now()
             
             if start_str:
                 start = datetime.strptime(start_str, "%Y-%m-%d")
             else:
                 start = now.replace(hour=0, minute=0, second=0, microsecond=0)
                 
             if end_str:
                 end = datetime.strptime(end_str, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
             else:
                 end = now.replace(hour=23, minute=59, second=59)

             date_label = f"{start.strftime('%b %d')} - {end.strftime('%b %d')}"
             if start.date() == now.date() and end.date() == now.date():
                 date_label = "Today's Data"
             
             # KPI: Total Visits
             total_visits = db.query(models.VisitHistory).filter(
                 models.VisitHistory.visit_date >= start,
                 models.VisitHistory.visit_date <= end
             ).count()

             # KPI: Avg Service Time
             avg_svc = db.query(func.avg(models.VisitHistory.duration_seconds)).filter(
                 models.VisitHistory.visit_date >= start,
                 models.VisitHistory.visit_date <= end,
                 models.VisitHistory.duration_seconds != None
             ).scalar()
             avg_service_time = int(avg_svc / 60) if avg_svc else 0

             # Chart: Visits by Room
             room_stats = db.query(models.VisitHistory.room, func.count(models.VisitHistory.id)).filter(
                 models.VisitHistory.visit_date >= start,
                 models.VisitHistory.visit_date <= end
             ).group_by(models.VisitHistory.room).all()
             room_data = {r: c for r, c in room_stats if r}
             
             # Table: Staff Performance (Visits & Avg Time)
             staff_stats = db.query(
                 models.User.full_name, 
                 func.count(models.VisitHistory.id),
                 func.avg(models.VisitHistory.duration_seconds)
             ).join(models.User, models.VisitHistory.doctor_id == models.User.id)\
              .filter(
                  models.VisitHistory.visit_date >= start,
                  models.VisitHistory.visit_date <= end
              )\
              .group_by(models.User.full_name).all()
             
             staff_data = []
             for name, count, avg_sec in staff_stats:
                 avg_min = int(avg_sec / 60) if avg_sec else 0
                 staff_data.append({"name": name, "visits": count, "avg_time": avg_min})

             # Chart: Hourly Traffic (7 AM - 10 PM)
             hourly_data = {f"{h}:00": 0 for h in range(7, 23)} 
             visits = db.query(models.VisitHistory.visit_date).filter(
                 models.VisitHistory.visit_date >= start,
                 models.VisitHistory.visit_date <= end
             ).all()
             
             for v in visits:
                  if v.visit_date:
                       h = v.visit_date.hour
                       if 7 <= h <= 22:
                            hourly_data[f"{h}:00"] += 1
             
        return await self.templates.TemplateResponse(request, "admin/analytics.html", context={
            "request": request,
            "avg_wait_time": 0, 
            "avg_service_time": avg_service_time,
            "total_visits": total_visits,
            "room_data": room_data,
            "staff_data": staff_data,
            "hourly_data": hourly_data,
            "date_label": date_label,
            "start_date": start.strftime("%Y-%m-%d"),
            "end_date": end.strftime("%Y-%m-%d")
        })


class DoctorsView(BaseView):
    name = "Doctors"
    icon = "fa-solid fa-user-doctor"
    category = "Hospital Resources"

    @expose("/doctors", methods=["GET"])
    async def doctors(self, request: Request):
        DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        with database.SessionLocal() as db:
            from backend.roster.models import Shift
            shifts = db.query(Shift).all()
            
            # Load all active departments that have doctors (excluding specific ones)
            excluded_departments = ["Administration", "Physiotherapy", "Tabara"]
            departments = db.query(models.Department).filter(
                ~models.Department.name.in_(excluded_departments)
            ).order_by(models.Department.name).all()
            
            dept_data = []
            for dept in departments:
                # Doctors in this department
                doctors = db.query(models.User).join(models.Role).filter(
                    models.Role.name == "Doctor",
                    models.User.department_id == dept.id,
                    models.User.is_active == True
                ).all()
                
                if not doctors:
                    continue
                    
                doctor_list = []
                for doc in doctors:
                    # Load existing roster entries
                    roster_entries = db.query(models.DoctorRoster).filter(
                        models.DoctorRoster.doctor_id == doc.id
                    ).all()
                    roster_map = {r.day_of_week: r.status for r in roster_entries}
                    
                    # Build full week, filling defaults
                    roster = []
                    for day in DAYS:
                        roster.append({
                            "day": day,
                            "day_short": day[:3],
                            "status": roster_map.get(day, "available")
                        })
                    
                    doctor_list.append({
                        "id": doc.id,
                        "name": doc.full_name or doc.username,
                        "username": doc.username,
                        "room": doc.room_number or "—",
                        "is_available": doc.is_available,
                        "roster": roster
                    })
                
                dept_data.append({
                    "id": dept.id,
                    "name": dept.name,
                    "doctors": doctor_list
                })

        return await self.templates.TemplateResponse(request, "admin/doctors.html", context={
            "request": request,
            "departments": dept_data,
            "days": DAYS,
            "shifts": shifts,
        })

    @expose("/doctors/update-roster", methods=["POST"])
    async def update_roster(self, request: Request):
        from starlette.responses import JSONResponse
        data = await request.json()
        doctor_id = data.get("doctor_id")
        day = data.get("day")
        new_status = str(data.get("status"))

        if not doctor_id or not day:
            return JSONResponse({"error": "Invalid input"}, status_code=400)

        with database.SessionLocal() as db:
            from backend.roster.models import Shift
            shifts = db.query(Shift).all()
            valid_statuses = ["available", "not_available"] + [str(s.id) for s in shifts]
            
            if new_status not in valid_statuses:
                return JSONResponse({"error": "Invalid shift or status"}, status_code=400)

            entry = db.query(models.DoctorRoster).filter(
                models.DoctorRoster.doctor_id == doctor_id,
                models.DoctorRoster.day_of_week == day
            ).first()

            if entry:
                entry.status = new_status
                entry.updated_at = datetime.utcnow()
            else:
                entry = models.DoctorRoster(
                    doctor_id=doctor_id,
                    day_of_week=day,
                    status=new_status,
                    updated_at=datetime.utcnow()
                )
                db.add(entry)
            db.commit()

        return JSONResponse({"success": True, "doctor_id": doctor_id, "day": day, "status": new_status})


# ─────────────────────────────────────────────────────────────
# Communications Admin Views
# ─────────────────────────────────────────────────────────────

class SMSHistoryAdmin(ModelView, model=models.SMSHistory):
    name = "SMS History"
    name_plural = "SMS History"
    icon = "fa-solid fa-comment-sms"
    category = "Communications"
    column_list = [
        models.SMSHistory.sent_at,
        models.SMSHistory.phone_number,
        "patient",
        models.SMSHistory.message_type,
        models.SMSHistory.status,
        "sent_by",
    ]
    column_sortable_list = [models.SMSHistory.sent_at, models.SMSHistory.status, models.SMSHistory.message_type]
    column_searchable_list = [models.SMSHistory.phone_number, models.SMSHistory.message_type]
    can_create = False
    can_edit = False
    can_delete = False
    can_view_details = True
    can_export = True

class SMSTemplateAdmin(ModelView, model=models.SMSTemplate):
    name = "SMS Templates"
    name_plural = "SMS Templates"
    icon = "fa-solid fa-file-lines"
    category = "Communications"
    column_list = [models.SMSTemplate.type, models.SMSTemplate.description, models.SMSTemplate.template]
    column_searchable_list = [models.SMSTemplate.type, models.SMSTemplate.description]
    can_create = True
    can_edit = True
    can_delete = True
    can_view_details = True
    can_export = True

# ─────────────────────────────────────────────────────────────
# System Admin View
# ─────────────────────────────────────────────────────────────

class SettingAdmin(ModelView, model=models.Setting):
    name = "System Settings"
    name_plural = "System Settings"
    icon = "fa-solid fa-gear"
    category = "System"
    column_list = [models.Setting.key, models.Setting.value, models.Setting.description, models.Setting.updated_at]
    column_sortable_list = [models.Setting.key, models.Setting.updated_at]
    column_searchable_list = [models.Setting.key, models.Setting.description]
    can_create = False
    can_edit = True
    can_delete = False
    can_view_details = True
    can_export = True

# ─────────────────────────────────────────────────────────────
# File Hub Admin Views
# ─────────────────────────────────────────────────────────────

class FileCategoryAdmin(ModelView, model=models.FileCategory):
    name = "File Categories"
    name_plural = "File Categories"
    icon = "fa-solid fa-folder"
    category = "File Hub"
    column_list = [models.FileCategory.id, models.FileCategory.name, models.FileCategory.description]
    column_searchable_list = [models.FileCategory.name]
    can_create = True
    can_edit = True
    can_delete = True
    can_view_details = True
    can_export = True

class DocumentAdmin(ModelView, model=models.Document):
    name = "Documents"
    name_plural = "Documents"
    icon = "fa-solid fa-file"
    category = "File Hub"
    column_list = [
        models.Document.original_name,
        "category",
        "uploaded_by",
        models.Document.upload_date,
        models.Document.file_size,
        models.Document.mime_type,
        models.Document.is_active,
    ]
    column_sortable_list = [models.Document.upload_date, models.Document.file_size, models.Document.original_name]
    column_searchable_list = [models.Document.original_name, models.Document.mime_type]
    can_create = False
    can_edit = True
    can_delete = True
    can_view_details = True
    can_export = True
    form_columns = [models.Document.category, models.Document.is_active]

class FileAuditLogAdmin(ModelView, model=models.FileAuditLog):
    name = "File Audit Log"
    name_plural = "File Audit Logs"
    icon = "fa-solid fa-shield-halved"
    category = "File Hub"
    column_list = [
        models.FileAuditLog.timestamp,
        "user",
        "document",
        models.FileAuditLog.action,
        models.FileAuditLog.ip_address,
    ]
    column_sortable_list = [models.FileAuditLog.timestamp, models.FileAuditLog.action]
    column_searchable_list = [models.FileAuditLog.action, models.FileAuditLog.ip_address]
    can_create = False
    can_edit = False
    can_delete = False
    can_view_details = True
    can_export = True

# ─────────────────────────────────────────────────────────────
# Roster Admin Views
# ─────────────────────────────────────────────────────────────

from backend.roster.models import (
    Unit, Shift, RosterDay, RosterAssignment, StaffAvailability,
    AuditLog as RosterAuditLog,
)

class DoctorRosterAdmin(ModelView, model=models.DoctorRoster):
    name = "Doctor Weekly Roster"
    name_plural = "Doctor Weekly Roster"
    icon = "fa-solid fa-calendar-check"
    category = "Roster"
    column_list = [
        "doctor",
        models.DoctorRoster.day_of_week,
        models.DoctorRoster.status,
        models.DoctorRoster.updated_at,
    ]
    column_sortable_list = [models.DoctorRoster.day_of_week, models.DoctorRoster.updated_at]
    column_searchable_list = [models.DoctorRoster.day_of_week, models.DoctorRoster.status]
    can_create = True
    can_edit = True
    can_delete = True
    can_export = True

class UnitAdmin(ModelView, model=Unit):
    name = "Units"
    name_plural = "Units"
    icon = "fa-solid fa-sitemap"
    category = "Roster"
    column_list = [Unit.id, Unit.name, "department"]
    column_searchable_list = [Unit.name]
    can_create = True
    can_edit = True
    can_delete = True
    can_export = True

class ShiftAdmin(ModelView, model=Shift):
    name = "Shifts"
    name_plural = "Shifts"
    icon = "fa-solid fa-clock"
    category = "Roster"
    column_list = [Shift.id, Shift.name, Shift.start_time, Shift.end_time]
    column_sortable_list = [Shift.name, Shift.start_time]
    column_searchable_list = [Shift.name]
    can_create = True
    can_edit = True
    can_delete = True
    can_export = True

class RosterDayAdmin(ModelView, model=RosterDay):
    name = "Roster Days"
    name_plural = "Roster Days"
    icon = "fa-solid fa-calendar-days"
    category = "Roster"
    column_list = [RosterDay.id, RosterDay.date, "creator", RosterDay.status, RosterDay.notes]
    column_sortable_list = [RosterDay.date, RosterDay.status]
    column_searchable_list = [RosterDay.status]
    can_create = True
    can_edit = True
    can_delete = True
    can_export = True
    can_view_details = True

class RosterAssignmentAdmin(ModelView, model=RosterAssignment):
    name = "Roster Assignments"
    name_plural = "Roster Assignments"
    icon = "fa-solid fa-user-clock"
    category = "Roster"
    column_list = [
        "roster_day", "staff", "department", "unit",
        RosterAssignment.shift_label,
        RosterAssignment.shift_start_time,
        RosterAssignment.shift_end_time,
        RosterAssignment.room_number,
        RosterAssignment.phone,
    ]
    column_sortable_list = [RosterAssignment.shift_label, RosterAssignment.created_at]
    column_searchable_list = [RosterAssignment.shift_label, RosterAssignment.phone, RosterAssignment.room_number]
    can_create = True
    can_edit = True
    can_delete = True
    can_export = True
    can_view_details = True

class StaffAvailabilityAdmin(ModelView, model=StaffAvailability):
    name = "Staff Availability"
    name_plural = "Staff Availability"
    icon = "fa-solid fa-user-check"
    category = "Roster"
    column_list = ["staff", StaffAvailability.date, StaffAvailability.available, StaffAvailability.reason]
    column_sortable_list = [StaffAvailability.date, StaffAvailability.available]
    can_create = True
    can_edit = True
    can_delete = True
    can_export = True

class RosterAuditLogAdmin(ModelView, model=RosterAuditLog):
    name = "Roster Audit Log"
    name_plural = "Roster Audit Logs"
    icon = "fa-solid fa-list-check"
    category = "Roster"
    column_list = ["actor", RosterAuditLog.action, RosterAuditLog.created_at]
    column_sortable_list = [RosterAuditLog.created_at]
    column_searchable_list = [RosterAuditLog.action]
    can_create = False
    can_edit = False
    can_delete = False
    can_export = True

# ─────────────────────────────────────────────────────────────
# Register all views
# ─────────────────────────────────────────────────────────────

admin.add_view(AnalyticsView)
admin.add_view(DoctorsView)

# User Management
admin.add_view(UserAdmin)
admin.add_view(RoleAdmin)

# Hospital Resources
admin.add_view(DepartmentAdmin)
admin.add_view(RoomAdmin)

# Medical Records
admin.add_view(PatientAdmin)
admin.add_view(QueueAdmin)
admin.add_view(VisitHistoryAdmin)
admin.add_view(PriorityLevelAdmin)

# Communications
admin.add_view(SMSHistoryAdmin)
admin.add_view(SMSTemplateAdmin)

# System
admin.add_view(SettingAdmin)

# File Hub
admin.add_view(FileCategoryAdmin)
admin.add_view(DocumentAdmin)
admin.add_view(FileAuditLogAdmin)

# Roster
admin.add_view(DoctorRosterAdmin)
admin.add_view(UnitAdmin)
admin.add_view(ShiftAdmin)
admin.add_view(RosterDayAdmin)
admin.add_view(RosterAssignmentAdmin)
admin.add_view(StaffAvailabilityAdmin)
admin.add_view(RosterAuditLogAdmin)
admin.add_view(PatientVitalsAdmin)

# Include modules at bottom to avoid circular imports
from backend.roster.router import router as roster_router
app.include_router(roster_router, prefix="/roster", tags=["Roster"])

from backend.routers.files import router as files_router
app.include_router(files_router, prefix="/api/files", tags=["Files"])

# ==========================================
# Nursing Portal - Observation & Medication
# ==========================================

@app.post("/patients/{patient_id}/observation-notes", response_model=schemas.ObservationNoteResponse)
async def create_observation_note(
    patient_id: int, 
    note: schemas.ObservationNoteCreate, 
    db: Session = Depends(get_db)
):
    # Verify patient exists
    patient = db.query(models.Patient).filter(models.Patient.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
        
    db_note = models.ObservationNote(
        patient_id=patient_id,
        nurse_id=note.nurse_id,
        content=note.content
    )
    db.add(db_note)
    db.commit()
    db.refresh(db_note)
    return db_note

@app.get("/patients/{patient_id}/observation-notes", response_model=List[schemas.ObservationNoteResponse])
async def list_observation_notes(
    patient_id: int, 
    db: Session = Depends(get_db)
):
    return db.query(models.ObservationNote).filter(models.ObservationNote.patient_id == patient_id).order_by(models.ObservationNote.created_at.desc()).all()

@app.post("/patients/{patient_id}/medications", response_model=schemas.MedicationAdministrationResponse)
async def administer_medication(
    patient_id: int, 
    admin: schemas.MedicationAdministrationCreate, 
    db: Session = Depends(get_db)
):
    # Verify patient exists
    patient = db.query(models.Patient).filter(models.Patient.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
        
    db_admin = models.MedicationAdministration(
        patient_id=patient_id,
        nurse_id=admin.nurse_id,
        medication_name=admin.medication_name,
        dosage=admin.dosage,
        route=admin.route,
        notes=admin.notes
    )
    db.add(db_admin)
    db.commit()
    db.refresh(db_admin)
    return db_admin

@app.get("/patients/{patient_id}/medications", response_model=List[schemas.MedicationAdministrationResponse])
async def list_medications(
    patient_id: int, 
    db: Session = Depends(get_db)
):
    return db.query(models.MedicationAdministration).filter(models.MedicationAdministration.patient_id == patient_id).order_by(models.MedicationAdministration.administered_at.desc()).all()


# --- Vitals / Triage Endpoints ---

@app.post("/patients/{patient_id}/vitals", response_model=schemas.PatientVitalsResponse)
async def create_patient_vitals(
    patient_id: int,
    vitals: schemas.PatientVitalsCreate,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user_optional)
):
    """Record new vitals for a patient (Triage)"""
    # Verify patient exists
    patient = db.query(models.Patient).filter(models.Patient.id == patient_id).first()
    if not patient:
        raise HTTPException(status_code=404, detail="Patient not found")
        
    # Use nurse_id from request if authenticated user is not present (simplified portal access)
    n_id = current_user.id if current_user else vitals.nurse_id
    
    db_vitals = models.PatientVitals(
        **vitals.dict(exclude={'nurse_id'}),
        nurse_id=n_id
    )
    db.add(db_vitals)
    db.commit()
    db.refresh(db_vitals)
    return db_vitals

@app.get("/patients/{patient_id}/vitals", response_model=List[schemas.PatientVitalsResponse])
async def list_patient_vitals(
    patient_id: int,
    db: Session = Depends(get_db)
):
    """List historical vitals for a patient"""
    return db.query(models.PatientVitals).filter(
        models.PatientVitals.patient_id == patient_id
    ).order_by(models.PatientVitals.recorded_at.desc()).all()


# ==========================================
# Nursing Portal - Settings & Profile
# ==========================================

@app.put("/users/{user_id}/room", status_code=200)
async def update_nurse_room(
    user_id: int,
    room_data: schemas.UserRoomUpdate,
    db: Session = Depends(get_db)
):
    """Updates a nurse's room assignment (unprotected for demonstration)."""
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
        
    user.room_number = room_data.room_number
    db.commit()
    db.refresh(user)
    return {"message": "Room updated", "room_number": user.room_number}

# ==========================================
# Admin - Change User Picture
# ==========================================

@app.post("/users/{user_id}/profile-picture")
async def upload_user_picture(
    user_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Allows an administrator to change a staff member's profile picture."""
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    # Storage setup
    upload_dir = "backend/static/uploads/profiles"
    os.makedirs(upload_dir, exist_ok=True)

    # Unique file persistence
    ext = os.path.splitext(file.filename)[1]
    filename = f"user_{user_id}_{uuid.uuid4().hex}{ext}"
    file_path = os.path.join(upload_dir, filename)

    # Physical storage
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Database update
    # Note: Using absolute URL for easy frontend rendering
    base_url = "https://localhost:8000"
    profile_url = f"{base_url}/static/uploads/profiles/{filename}"
    user.profile_picture = profile_url
    db.commit()
    db.refresh(user)

    return {"message": "Profile picture updated successfully", "profile_picture": profile_url}
