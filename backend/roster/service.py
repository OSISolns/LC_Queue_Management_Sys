import math
import io
import re as _re
from typing import List, Dict, Optional
from datetime import date, time
from sqlalchemy.orm import Session
from backend.roster import repository, schemas, models
from backend.models import User as ClinicUser, Department
from collections import defaultdict


def _parse_time(value: str) -> Optional[time]:
    """Parse a time string like '08:00', '8:00 AM', or a decimal (Excel serial)."""
    if not value:
        return None
    value = str(value).strip()
    # Handle Excel decimal time (e.g. 0.333 = 08:00)
    try:
        decimal = float(value)
        total_minutes = int(decimal * 24 * 60)
        return time(total_minutes // 60 % 24, total_minutes % 60)
    except ValueError:
        pass
    # Handle HH:MM or H:MM AM/PM
    import re
    m = re.match(r'(\d{1,2}):(\d{2})(?:\s*(AM|PM))?', value, re.IGNORECASE)
    if m:
        h, mn, meridiem = int(m.group(1)), int(m.group(2)), m.group(3)
        if meridiem:
            if meridiem.upper() == 'PM' and h != 12:
                h += 12
            elif meridiem.upper() == 'AM' and h == 12:
                h = 0
        return time(h % 24, mn)
    return None


def parse_roster_file(file_bytes: bytes, ext: str, db: Session, roster_day_id: int, user_id: int):
    """
    Parse a .docx or .xlsx file and create roster assignments.
    Now more robust to handle multiple header rows and automatic staff creation.
    """
    # Load lookup maps
    dept_map = {d.name.lower(): d for d in db.query(Department).all()}
    from backend.models import Role
    doctor_role = db.query(Role).filter(Role.name == "Doctor").first()
    default_role_id = doctor_role.id if doctor_role else 2

    # Index staff by BOTH full_name and username
    staff_map = {}
    def refresh_staff_map():
        nonlocal staff_map
        staff_map = {}
        for u in db.query(ClinicUser).filter(ClinicUser.is_active == True).all():
            if u.full_name:
                staff_map[u.full_name.lower()] = u
            if u.username:
                staff_map[u.username.lower()] = u

    refresh_staff_map()

    def find_dept(name: str):
        if not name: return dept_map.get("physiotherapy")
        key = str(name).strip().lower()
        if not key or key in ("0", "department 0"):
            return dept_map.get("physiotherapy")
        if key in dept_map:
            return dept_map[key]
        for k, d in dept_map.items():
            if key in k or k in key:
                return d
        return dept_map.get("physiotherapy")

    def get_or_create_staff(name: str, dept_id: Optional[int] = None):
        if not name or len(name.strip()) < 3: return None
        # Clean name: remove "Dr", "Dr.", "Mr", etc. prefixes for matching
        import re as _re
        clean_name = _re.sub(r'^(Dr\.?|Mr\.?|Mrs\.?|Ms\.?)\s+', '', name.strip(), flags=_re.IGNORECASE).strip()
        
        staff = staff_map.get(clean_name.lower())
        if not staff:
            # Try original name if clean didn't work
            staff = staff_map.get(name.strip().lower())
            
        if not staff:
            # Partial match search
            for k, u in staff_map.items():
                if clean_name.lower() in k or k in clean_name.lower():
                    return u

            # Create new staff if not found
            username = clean_name.lower().replace(" ", "_")
            # Ensure unique username
            base_username = username
            counter = 1
            while db.query(ClinicUser).filter(ClinicUser.username == username).first():
                username = f"{base_username}{counter}"
                counter += 1
            
            # Use hashed password "welcome123" as default
            from backend.main import get_password_hash
            new_user = ClinicUser(
                username=username,
                full_name=clean_name,
                hashed_password=get_password_hash("welcome123"),
                role_id=default_role_id,
                department_id=dept_id,
                is_active=True
            )
            db.add(new_user)
            db.commit()
            db.refresh(new_user)
            refresh_staff_map() # Update local cache
            return new_user
        return staff

    rows_raw = []
    if ext == '.docx':
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        if not doc.tables:
            raise ValueError("The uploaded .docx file contains no tables.")
        table = doc.tables[0]
        for row in table.rows:
            rows_raw.append([cell.text.strip() for cell in row.cells])
    elif ext in ('.xlsx', '.xls'):
        import openpyxl
        from datetime import time as dt_time
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        ws = wb.active
        for row in ws.iter_rows(values_only=True):
             # Keep time objects as they are, but strip strings
             processed_row = []
             for v in row:
                 if isinstance(v, dt_time):
                     processed_row.append(v)
                 else:
                     processed_row.append(str(v).strip() if v is not None else "")
             rows_raw.append(processed_row)
    
    if len(rows_raw) < 2:
        raise ValueError("File contains insufficient data.")

    # Find the header row (contains "Staff" or "Doctor" or "Department")
    header_idx = -1
    STAFF_KEYWORDS = ['staff', 'doctor', 'nurse', 'name', 'employee']
    for i, row in enumerate(rows_raw[:6]):
        if any(k in str(cell).lower() for cell in row for k in STAFF_KEYWORDS):
            header_idx = i
            break
    
    if header_idx == -1:
        header_idx = 0 # Fallback
        
    headers = [str(h).lower() for h in rows_raw[header_idx]]
    
    # Identify indices
    cols = {}
    STAFF_ALIASES  = ['staff name', 'staff', 'name', 'fullname', 'full name', 'employee', 'doctor', 'doctors', 'nurse']
    DEPT_ALIASES   = ['department', 'dept', 'department name', 'unit']
    ROLE_ALIASES   = ['role', 'position', 'designation', 'title']
    SHIFT_ALIASES  = ['shift', 'shift name', 'shift type', 'duty']
    
    # We might have multiple staff columns (e.g. "Morning", "Evening")
    # or a single "Staff" column + "Shift" column.
    staff_cols = []
    dept_idx = -1
    start_time_idx = -1
    end_time_idx = -1
    
    START_TIME_ALIASES = ['start time', 'start', 'from']
    END_TIME_ALIASES = ['end time', 'end', 'to']

    for i, h in enumerate(headers):
        if any(a in h for a in DEPT_ALIASES): dept_idx = i
        if any(a in h for a in STAFF_ALIASES): staff_cols.append(i)
        if any(a in h for a in START_TIME_ALIASES): start_time_idx = i
        if any(a in h for a in END_TIME_ALIASES): end_time_idx = i

    # If row below headers contains shift hints (Morning/Evening/...)
    SHIFT_KEYWORDS = ['morning', 'evening', 'night', 'afternoon', 'day', 'on-call', 'duty']
    shift_hints = {}
    if header_idx + 1 < len(rows_raw):
        potential_hints = rows_raw[header_idx + 1]
        is_hint_row = False
        for i in staff_cols:
            hint = str(potential_hints[i]).strip() if i < len(potential_hints) else ""
            if hint:
                # Only treat as a hint if it contains a shift keyword
                if any(k in hint.lower() for k in SHIFT_KEYWORDS):
                    # Clean hint: "Morning /Time" -> "Morning"
                    hint_clean = _re.split(r'[\/:\(]', hint)[0].strip()
                    shift_hints[i] = hint_clean
                    is_hint_row = True
        
        # If no keywords found, clear hints to avoid misidentification
        if not is_hint_row:
            shift_hints = {}

    # Clear old assignments
    repository.clear_assignments(db, roster_day_id)
    
    assignments = []
    errors = []
    
    def parse_cell_staff(cell_text: str):
        """Split cell by newline or common separators and extract time hints."""
        # Normalize: replace + or & with \n to treat them as separate lines
        cell_text = cell_text.replace('+', '\n').replace('&', '\n')
        lines = [l.strip() for l in cell_text.split('\n') if l.strip()]
        results = []
        for line in lines:
            # Extract time hint like (9am-4pm)
            time_match = _re.search(r'\((?:From\s+)?(\d{1,2}(?::\d{2})?\s*(?:am|pm)?)\s*-\s*(\d{1,2}(?::\d{2})?\s*(?:am|pm)?)\)', line, _re.IGNORECASE)
            start_t, end_t = time(8, 0), time(16, 0)
            if time_match:
                start_t = _parse_time(time_match.group(1)) or start_t
                end_t = _parse_time(time_match.group(2)) or end_t
            
            # Strip time and Role markers like (Ped) or (Gen) if they wrap a time
            clean_name = _re.sub(r'\s*\([^)]*\d+[^)]*\)', '', line).strip()
            # Also strip anything after Dr if it looks like a name but has trailing notes
            clean_name = _re.sub(r'\(.*?\)', '', clean_name).strip()
            
            if clean_name:
                results.append((clean_name, start_t, end_t))
        return results

    # Process data rows
    for r_idx in range(header_idx + 1, len(rows_raw)):
        row = rows_raw[r_idx]
        # Skip if it's the shift-hint row we already processed (if any)
        if any(h in str(cell) for cell in row for h in shift_hints.values()):
            if r_idx < header_idx + 3: continue 

        dept_val = row[dept_idx] if dept_idx != -1 and dept_idx < len(row) else ""
        dept = find_dept(dept_val)
        dept_id = dept.id if dept else 0

        for s_idx in staff_cols:
            if s_idx >= len(row): continue
            cell_content = row[s_idx]
            if not cell_content or len(cell_content) < 3: continue
            
            shift_label = shift_hints.get(s_idx, "Day")
            
            # Use specific time columns if they exist
            row_start = row[start_time_idx] if start_time_idx != -1 and start_time_idx < len(row) else None
            row_end = row[end_time_idx] if end_time_idx != -1 and end_time_idx < len(row) else None
            
            from datetime import time as dt_time
            def get_time(val, default):
                if isinstance(val, dt_time): return val
                if not val: return default
                return _parse_time(str(val)) or default

            staff_entries = parse_cell_staff(cell_content)
            for s_name, s_start, s_end in staff_entries:
                # Override with row-level times if available
                final_start = get_time(row_start, s_start)
                final_end = get_time(row_end, s_end)
                
                staff = get_or_create_staff(s_name, dept_id)
                if staff:
                    assignments.append(schemas.RosterAssignmentCreate(
                        roster_day_id=roster_day_id,
                        department_id=dept_id,
                        unit_id=None,
                        staff_id=staff.id,
                        shift_start_time=final_start,
                        shift_end_time=final_end,
                        shift_label=shift_label,
                        phone=staff.phone_number
                    ))
                else:
                    errors.append(f"Row {r_idx+1}: Could not find or create staff '{s_name}'")

    if not assignments:
        repository.log_audit_action(db, user_id, "UPLOAD_FAILED", {"file": "docx/xlsx", "errors": errors})
        return 0, errors

    repository.bulk_create_assignments(db, assignments)
    repository.log_audit_action(
        db, user_id, "UPLOAD_ROSTER",
        {"roster_day_id": roster_day_id, "assignments_created": len(assignments), "parse_errors": errors}
    )
    return len(assignments), errors




def generate_roster(
    db: Session,
    roster_day_id: int,
    request: schemas.RosterGenerateRequest,
    user_id: int
):
    roster_day = repository.get_roster_day(db, roster_day_id)
    if not roster_day:
        raise ValueError("Roster day not found")
        
    if roster_day.status == "published":
        raise ValueError("Cannot regenerate a published roster")

    target_date = roster_day.date

    # 1. Clear ALL existing assignments for this day before regenerating
    # This ensures no old dummy/stale data persists across generations
    repository.clear_assignments(db, roster_day_id)
    assigned_staff_ids = {}  # start fresh

    # 2. Fetch real clinic users — only include clinical staff with a department
    # Exclude administrative roles that shouldn't be on shift rosters
    excluded_roles = {"Admin", "SMS Officer", "Helpdesk"}
    unavail_staff_ids = set(repository.get_staff_unavailability(db, target_date))
    exclude_ids = set(request.exclude_staff_ids)

    clinic_users = db.query(ClinicUser).filter(ClinicUser.is_active == True).all()
    staff_list = [
        u for u in clinic_users
        if (u.role.category if u.role else "Staff") not in excluded_roles
        and u.id not in unavail_staff_ids
        and u.id not in exclude_ids
        and u.department_id is not None  # Must belong to a department
    ]

    # 3. Get active shifts
    shifts = repository.get_active_shifts(db)
    if not shifts:
        raise ValueError("No active shifts found in the system.")

    # Group staff by department
    dept_staff = defaultdict(list)
    for u in staff_list:
        dept_id = u.department_id or 0  # 0 = no department
        dept_staff[dept_id].append(u)

    new_assignments = []

    # Round-Robin assignment: 1 staff per shift per department
    for dept_id, staff in dept_staff.items():
        staff.sort(key=lambda u: u.id)  # deterministic ordering
        staff_idx = 0

        for shift in shifts:
            needed = request.min_staff_per_shift

            assigned_count = 0
            while assigned_count < needed:
                if staff_idx >= len(staff):
                    break  # Not enough staff for this dept

                candidate = staff[staff_idx]
                staff_idx += 1

                if candidate.id in assigned_staff_ids:
                    continue  # Already assigned today, skip

                assignment = schemas.RosterAssignmentCreate(
                    roster_day_id=roster_day_id,
                    department_id=dept_id or 0,
                    unit_id=None,
                    staff_id=candidate.id,
                    shift_start_time=shift.start_time,
                    shift_end_time=shift.end_time,
                    shift_label=shift.name,
                    phone=candidate.phone_number  # real Users field
                )
                new_assignments.append(assignment)
                assigned_staff_ids[candidate.id] = assignment
                assigned_count += 1

    # 4. Bulk insert
    repository.bulk_create_assignments(db, new_assignments)

    # Audit trail
    repository.log_audit_action(
        db, user_id, "GENERATE_ROSTER",
        {"roster_day_id": roster_day_id, "assignments_created": len(new_assignments)}
    )

    return get_roster_summary(db, roster_day_id)


def get_roster_summary(db: Session, roster_day_id: int):
    roster_day = repository.get_roster_day(db, roster_day_id)
    if not roster_day:
        return None

    assignments = repository.get_assignments_for_day(db, roster_day_id)

    # Pre-load real clinic users for name/role resolution
    clinic_users = {
        u.id: u for u in db.query(ClinicUser).filter(ClinicUser.is_active == True).all()
    }
    
    # Pre-load departments for name resolution
    departments_map = {
        d.id: d.name for d in db.query(Department).all()
    }

    # Group by Department -> Unit
    dep_map = {}
    
    for a in assignments:
        d_id = a.department_id
        if d_id not in dep_map:
            # Fallback for ID 0 or missing ID
            dept_name = departments_map.get(d_id)
            if not dept_name:
                if d_id == 0:
                    dept_name = "PHYSIOTHERAPY"
                else:
                    dept_name = f"Dept {d_id}"
            
            dep_map[d_id] = {
                "department_id": d_id,
                "department_name": dept_name,
                "units": {}
            }
        
        u_id = a.unit_id or 0
        if u_id not in dep_map[d_id]["units"]:
            dep_map[d_id]["units"][u_id] = {
                "unit_id": a.unit_id,
                "unit_name": "General",
                "assignments": []
            }

        # Resolve staff name/role from real Users table
        user = clinic_users.get(a.staff_id)
        staff_name = (user.full_name or user.username) if user else f"Staff {a.staff_id}"
        staff_role = (user.role.name if user and user.role else "Staff")
            
        dep_map[d_id]["units"][u_id]["assignments"].append({
            "id": a.id,
            "staff_id": a.staff_id,
            "staff_name": staff_name,
            "role": staff_role,
            "shift_start_time": a.shift_start_time,
            "shift_end_time": a.shift_end_time,
            "shift_label": a.shift_label,
            "phone": a.phone,
            "room_number": a.room_number or (user.room_number if user else None)
        })

    # Convert to schema shape
    departments = []
    for d_id, d_data in dep_map.items():
        units_list = []
        for u_id, u_data in d_data["units"].items():
            units_list.append(schemas.UnitRosterGroup(**u_data))
        d_data["units"] = units_list
        departments.append(schemas.DepartmentRosterGroup(**d_data))

    return schemas.RosterDayDetailResponse(
        id=roster_day.id,
        date=roster_day.date,
        notes=roster_day.notes,
        created_by_user_id=roster_day.created_by_user_id,
        status=roster_day.status,
        departments=departments
    )


def get_hourly_coverage(db: Session, target_date: date):
    roster_day = repository.get_roster_day_by_date(db, target_date)
    if not roster_day or roster_day.status != "published":
        return None # or empty

    assignments = repository.get_assignments_for_day(db, roster_day.id)
    
    # department -> hour -> count
    coverage = defaultdict(lambda: defaultdict(int))
    dept_names = {}
    
    for a in assignments:
        dept_id = a.department_id
        dept_names[dept_id] = a.department.name if a.department else f"Dept {dept_id}"
        
        start_hour = a.shift_start_time.hour
        end_hour = a.shift_end_time.hour
        
        # handle overnight shifts? assume same day for MVP
        for hr in range(start_hour, end_hour):
            coverage[dept_id][hr] += 1
            
    response = []
    for d_id, hours_map in coverage.items():
        hr_list = [schemas.HourlyCoverage(hour=hr, active_staff_count=cnt) for hr, cnt in hours_map.items()]
        response.append(schemas.DepartmentCoverage(
            department_id=d_id,
            department_name=dept_names[d_id],
            hourly_coverage=hr_list
        ))
        
    return schemas.CoverageResponse(date=target_date, departments=response)
